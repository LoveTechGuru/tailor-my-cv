# ============================================================
#  TailorMyCV — app.py
#  Micro-SaaS CV Tailoring Tool for the Indian Market
#  Built with Streamlit + Claude (Anthropic) + Razorpay
# ============================================================

import streamlit as st
import anthropic
import io
import os
import json
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Optional parsers (graceful fallback) ────────────────────
try:
    import PyPDF2
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    import docx2txt
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


# ════════════════════════════════════════════════════════════
#  PAGE CONFIG  (must be first Streamlit call)
# ════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="TailorMyCV — AI-Powered Resume Tailoring",
    page_icon="✦",
    layout="centered",
    initial_sidebar_state="collapsed",
)


# ════════════════════════════════════════════════════════════
#  GLOBAL CSS — Premium Indigo/Cream SaaS Theme
# ════════════════════════════════════════════════════════════
st.markdown("""
<style>
/* ── Google Fonts ── */
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@300;400;600;700&family=DM+Mono:wght@400;500&display=swap');

/* ── CSS Variables ── */
:root {
  --indigo:     #4F46E5;
  --indigo-dk:  #3730A3;
  --indigo-lt:  #EEF2FF;
  --accent:     #F59E0B;
  --cream:      #FAFAF7;
  --dark:       #1E1B4B;
  --muted:      #6B7280;
  --border:     #E5E7EB;
  --success:    #10B981;
  --radius:     14px;
  --shadow:     0 4px 24px rgba(79,70,229,0.10);
}

/* ── Base Reset ── */
html, body, [class*="css"] {
  font-family: 'Sora', sans-serif !important;
  background: var(--cream) !important;
  color: var(--dark) !important;
}

/* ── Hide Streamlit Chrome ── */
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 1.5rem 1rem 4rem !important; max-width: 780px !important; }

/* ── Masthead ── */
.masthead {
  text-align: center;
  padding: 2.5rem 1rem 1.5rem;
  background: linear-gradient(135deg, var(--indigo) 0%, var(--indigo-dk) 100%);
  border-radius: var(--radius);
  margin-bottom: 2rem;
  box-shadow: var(--shadow);
}
.masthead-logo {
  font-size: 2rem;
  letter-spacing: -1px;
  color: #fff;
  font-weight: 700;
}
.masthead-logo span { color: var(--accent); }
.masthead-sub {
  color: rgba(255,255,255,0.75);
  font-size: 0.88rem;
  margin-top: 0.35rem;
  letter-spacing: 0.05em;
}

/* ── Step Progress Bar ── */
.step-bar {
  display: flex;
  align-items: center;
  justify-content: center;
  gap: 0;
  margin-bottom: 2rem;
}
.step-pill {
  display: flex;
  align-items: center;
  gap: 6px;
  padding: 6px 16px;
  border-radius: 999px;
  font-size: 0.78rem;
  font-weight: 600;
  color: var(--muted);
  background: var(--border);
  transition: all .3s;
}
.step-pill.active {
  background: var(--indigo);
  color: #fff;
  box-shadow: 0 2px 12px rgba(79,70,229,0.35);
}
.step-pill.done {
  background: var(--success);
  color: #fff;
}
.step-connector {
  width: 28px;
  height: 2px;
  background: var(--border);
}

/* ── Card ── */
.card {
  background: #fff;
  border-radius: var(--radius);
  border: 1px solid var(--border);
  padding: 1.8rem 1.5rem;
  margin-bottom: 1.2rem;
  box-shadow: var(--shadow);
}
.card-title {
  font-size: 1.05rem;
  font-weight: 700;
  color: var(--indigo-dk);
  margin-bottom: 0.8rem;
}

/* ── Mode Cards (radio replacement) ── */
.mode-grid {
  display: grid;
  grid-template-columns: repeat(3, 1fr);
  gap: 10px;
  margin-top: 0.5rem;
}
@media(max-width:600px){ .mode-grid { grid-template-columns: 1fr; } }
.mode-card {
  border: 2px solid var(--border);
  border-radius: var(--radius);
  padding: 14px 12px;
  cursor: pointer;
  transition: all .2s;
  text-align: center;
  background: #fff;
}
.mode-card:hover { border-color: var(--indigo); background: var(--indigo-lt); }
.mode-card.selected { border-color: var(--indigo); background: var(--indigo-lt); }
.mode-icon { font-size: 1.6rem; }
.mode-label { font-size: 0.8rem; font-weight: 700; color: var(--dark); margin-top: 6px; }
.mode-desc { font-size: 0.72rem; color: var(--muted); margin-top: 3px; }

/* ── Streamlit Widgets Override ── */
.stTextArea textarea {
  border-radius: 10px !important;
  border: 1.5px solid var(--border) !important;
  font-family: 'DM Mono', monospace !important;
  font-size: 0.82rem !important;
  background: #FAFAFA !important;
}
.stTextArea textarea:focus {
  border-color: var(--indigo) !important;
  box-shadow: 0 0 0 3px rgba(79,70,229,0.12) !important;
}
.stFileUploader {
  border: 2px dashed var(--indigo) !important;
  border-radius: var(--radius) !important;
  background: var(--indigo-lt) !important;
  padding: 1rem !important;
}
div[data-testid="stFileUploadDropzone"] {
  background: var(--indigo-lt) !important;
}

/* ── Buttons ── */
.stButton > button {
  background: var(--indigo) !important;
  color: #fff !important;
  border: none !important;
  border-radius: 999px !important;
  padding: 0.6rem 2rem !important;
  font-family: 'Sora', sans-serif !important;
  font-weight: 600 !important;
  font-size: 0.9rem !important;
  letter-spacing: 0.02em !important;
  transition: all .2s !important;
  box-shadow: 0 3px 14px rgba(79,70,229,0.30) !important;
  width: 100% !important;
}
.stButton > button:hover {
  background: var(--indigo-dk) !important;
  transform: translateY(-1px) !important;
  box-shadow: 0 6px 20px rgba(79,70,229,0.38) !important;
}

/* ── Pay Button ── */
.pay-btn-wrap { text-align: center; margin: 1rem 0; }
.pay-btn {
  display: inline-block;
  background: linear-gradient(135deg, #F59E0B 0%, #D97706 100%);
  color: #fff !important;
  font-family: 'Sora', sans-serif;
  font-weight: 700;
  font-size: 1rem;
  padding: 14px 40px;
  border-radius: 999px;
  text-decoration: none !important;
  box-shadow: 0 4px 18px rgba(245,158,11,0.40);
  transition: all .2s;
  border: none;
  cursor: pointer;
}
.pay-btn:hover { transform: translateY(-2px); box-shadow: 0 8px 24px rgba(245,158,11,0.50); }

/* ── Badges ── */
.badge {
  display: inline-block;
  background: var(--indigo-lt);
  color: var(--indigo);
  font-size: 0.72rem;
  font-weight: 700;
  padding: 3px 10px;
  border-radius: 999px;
  margin-right: 6px;
  margin-bottom: 4px;
}
.badge-success {
  background: #D1FAE5;
  color: #065F46;
}

/* ── Result Preview ── */
.result-header {
  display: flex;
  align-items: center;
  justify-content: space-between;
  margin-bottom: 1rem;
}
.verified-pill {
  background: #D1FAE5;
  color: #065F46;
  border-radius: 999px;
  font-size: 0.75rem;
  font-weight: 700;
  padding: 4px 12px;
}

/* ── Info / Warning boxes ── */
.info-box {
  background: var(--indigo-lt);
  border-left: 4px solid var(--indigo);
  border-radius: 8px;
  padding: 12px 16px;
  font-size: 0.83rem;
  color: var(--indigo-dk);
  margin: 0.8rem 0;
}
.warn-box {
  background: #FFFBEB;
  border-left: 4px solid var(--accent);
  border-radius: 8px;
  padding: 12px 16px;
  font-size: 0.83rem;
  color: #92400E;
  margin: 0.8rem 0;
}
.success-box {
  background: #D1FAE5;
  border-left: 4px solid var(--success);
  border-radius: 8px;
  padding: 12px 16px;
  font-size: 0.83rem;
  color: #065F46;
  margin: 0.8rem 0;
}

/* ── Footer ── */
.footer {
  text-align: center;
  font-size: 0.74rem;
  color: var(--muted);
  margin-top: 3rem;
  padding-top: 1rem;
  border-top: 1px solid var(--border);
}
</style>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════
#  SESSION STATE INIT
# ════════════════════════════════════════════════════════════
DEFAULTS = {
    "step": 1,
    "cv_text": "",
    "jd_text": "",
    "mode": "ATS-Optimizer",
    "tailored_cv": "",
    "payment_verified": True,
    "pay_order_id": "",
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ════════════════════════════════════════════════════════════
#  HELPERS
# ════════════════════════════════════════════════════════════
def extract_text_from_pdf(file_bytes: bytes) -> str:
    if not PDF_OK:
        return ""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    return "\n".join(
        page.extract_text() or "" for page in reader.pages
    )


def extract_text_from_docx(file_bytes: bytes) -> str:
    if not DOCX_OK:
        return ""
    return docx2txt.process(io.BytesIO(file_bytes))


def get_anthropic_client():
    try:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception:
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        st.error("⚠️ ANTHROPIC_API_KEY not set. Add it in Streamlit Secrets.")
        st.stop()
    return anthropic.Anthropic(api_key=api_key)


MODE_META = {
    "ATS-Optimizer": {
        "icon": "🎯",
        "desc": "Beat ATS filters with keyword matching",
        "system_add": (
            "Focus on ATS keyword optimization. Mirror the exact terminology from the "
            "job description. Use standard section headings (Experience, Skills, Education). "
            "Quantify every achievement with numbers where possible."
        ),
    },
    "Executive Tone": {
        "icon": "💼",
        "desc": "Senior leadership & boardroom gravitas",
        "system_add": (
            "Rewrite with a commanding executive tone suitable for C-suite and senior director roles. "
            "Lead with strategic impact and business outcomes. Eliminate junior-level language. "
            "Every bullet must demonstrate leadership, ownership, or P&L impact."
        ),
    },
    "Career Switch": {
        "icon": "🔄",
        "desc": "Pivot your story to a new industry",
        "system_add": (
            "The candidate is making a career switch. Aggressively reframe transferable skills "
            "to align with the target role. Write a compelling Summary that bridges the gap. "
            "Downplay non-relevant experience and amplify any crossover skills."
        ),
    },
}


def build_system_prompt(mode: str) -> str:
    mode_instruction = MODE_META[mode]["system_add"]
    return f"""You are a world-class executive recruiter and resume strategist with 20 years of experience
at top-tier firms. You have placed candidates at Google, McKinsey, Goldman Sachs, and Fortune 500 companies.

Your task: Rewrite the candidate's CV to perfectly match the provided Job Description.

Core directives:
1. {mode_instruction}
2. NEVER invent experience or credentials — only reframe existing content.
3. Use strong action verbs: spearheaded, architected, orchestrated, accelerated, transformed.
4. Structure: Professional Summary → Key Skills → Work Experience → Education → Certifications.
5. Every work-experience bullet must follow: [Action Verb] + [Task] + [Result/Impact].
6. The output MUST be clean plain text, ready to be placed in a Word document.
7. Do NOT add any commentary, preamble, or markdown headers — output the CV content only.
8. Keep it to one page worth of content unless the role genuinely requires more.

Output the tailored CV now."""


def call_claude(cv_text: str, jd_text: str, mode: str) -> str:
    client = get_anthropic_client()
    user_msg = (
        f"=== ORIGINAL CV ===\n{cv_text}\n\n"
        f"=== JOB DESCRIPTION ===\n{jd_text}\n\n"
        f"Please tailor the CV above to match this job description using the {mode} mode."
    )
    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=2048,
        system=build_system_prompt(mode),
        messages=[{"role": "user", "content": user_msg}],
    )
    return message.content[0].text


def generate_docx(cv_text: str, mode: str) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # Watermark / branding line
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run(f"✦ TailorMyCV — {mode} Edition")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_paragraph("")  # spacer

    # Parse and write CV text
    for line in cv_text.strip().split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # Detect section headings (ALL CAPS or title-like short lines)
        if line.isupper() or (len(line) < 40 and line.endswith(":")):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
        elif line.startswith("•") or line.startswith("-"):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line.lstrip("•- ").strip())
            run.font.size = Pt(10)
        else:
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(10) if p.runs else None

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def get_razorpay_link() -> str:
    """Return Razorpay payment link from secrets or env."""
    try:
        return st.secrets["RAZORPAY_PAYMENT_LINK"]
    except Exception:
        return os.environ.get(
            "RAZORPAY_PAYMENT_LINK",
            "https://rzp.io/l/tailormycv-demo"   # ← Replace with your real link
        )


# ════════════════════════════════════════════════════════════
#  UI COMPONENTS
# ════════════════════════════════════════════════════════════
def render_masthead():
    st.markdown("""
    <div class="masthead">
      <div class="masthead-logo">Tailor<span>My</span>CV ✦</div>
      <div class="masthead-sub">AI-Powered Resume Tailoring · ₹20 per download · Instant results</div>
    </div>
    """, unsafe_allow_html=True)


def render_step_bar(current: int):
    steps = ["Upload CV", "Job Description", "Choose Mode", "Result & Pay"]
    pills = ""
    for i, label in enumerate(steps, 1):
        if i < current:
            cls = "done"
            icon = "✓"
        elif i == current:
            cls = "active"
            icon = str(i)
        else:
            cls = ""
            icon = str(i)
        pills += f'<div class="step-pill {cls}">{icon} {label}</div>'
        if i < len(steps):
            pills += '<div class="step-connector"></div>'
    st.markdown(f'<div class="step-bar">{pills}</div>', unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════
#  STEP PAGES
# ════════════════════════════════════════════════════════════

# ── STEP 1: Upload CV ──────────────────────────────────────
def page_step1():
    st.markdown('<div class="card"><div class="card-title">📄 Upload or Paste Your CV</div>', unsafe_allow_html=True)
    st.markdown('<div class="info-box">Supports PDF and DOCX uploads, or simply paste your CV text below.</div>', unsafe_allow_html=True)

    uploaded = st.file_uploader(
        "Upload CV (PDF or DOCX)",
        type=["pdf", "docx"],
        help="Max 5 MB",
        label_visibility="collapsed",
    )

    if uploaded:
        raw = uploaded.read()
        if uploaded.name.endswith(".pdf"):
            extracted = extract_text_from_pdf(raw)
        else:
            extracted = extract_text_from_docx(raw)

        if extracted.strip():
            st.session_state.cv_text = extracted
            st.success(f"✅ Extracted {len(extracted.split())} words from **{uploaded.name}**")
        else:
            st.warning("Could not auto-extract text. Please paste manually below.")

    pasted = st.text_area(
        "Or paste your CV text here:",
        value=st.session_state.cv_text,
        height=260,
        placeholder="Paste your full CV / resume here…",
        label_visibility="collapsed",
    )
    if pasted:
        st.session_state.cv_text = pasted

    st.markdown('</div>', unsafe_allow_html=True)

    if st.button("Continue → Job Description", use_container_width=True):
        if len(st.session_state.cv_text.strip()) < 50:
            st.error("Please provide your CV (at least 50 characters).")
        else:
            st.session_state.step = 2
            st.rerun()


# ── STEP 2: Job Description ────────────────────────────────
def page_step2():
    st.markdown('<div class="card"><div class="card-title">🎯 Paste the Job Description</div>', unsafe_allow_html=True)
    st.markdown('<div class="info-box">Copy the full JD from LinkedIn, Naukri, or any job portal.</div>', unsafe_allow_html=True)

    jd = st.text_area(
        "Job Description:",
        value=st.session_state.jd_text,
        height=280,
        placeholder="Paste the complete job description here…",
        label_visibility="collapsed",
    )
    st.session_state.jd_text = jd
    st.markdown('</div>', unsafe_allow_html=True)

    col1, col2 = st.columns([1, 2])
    with col1:
        if st.button("← Back"):
            st.session_state.step = 1
            st.rerun()
    with col2:
        if st.button("Continue → Select Mode", use_container_width=True):
            if len(st.session_state.jd_text.strip()) < 50:
                st.error("Please paste the job description (at least 50 characters).")
            else:
                st.session_state.step = 3
                st.rerun()


# ── STEP 3: Tailoring Mode ─────────────────────────────────
def page_step3():
    st.markdown('<div class="card"><div class="card-title">⚙️ Choose Your Tailoring Mode</div>', unsafe_allow_html=True)

    # Mode selector using columns
    cols = st.columns(3)
    for idx, (mode_key, meta) in enumerate(MODE_META.items()):
        with cols[idx]:
            selected = st.session_state.mode == mode_key
            border_style = "border: 2px solid #4F46E5; background: #EEF2FF;" if selected else "border: 2px solid #E5E7EB;"
            st.markdown(f"""
            <div class="mode-card {'selected' if selected else ''}" style="{border_style}">
              <div class="mode-icon">{meta['icon']}</div>
              <div class="mode-label">{mode_key}</div>
              <div class="mode-desc">{meta['desc']}</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button(f"Select", key=f"mode_{mode_key}", use_container_width=True):
                st.session_state.mode = mode_key
                st.rerun()

    st.markdown(f"""
    <div class="success-box" style="margin-top:1rem;">
      ✦ Selected: <strong>{st.session_state.mode}</strong> — {MODE_META[st.session_state.mode]['desc']}
    </div>
    """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    col1, col2 = st.columns([1, 2])
    with col1:
        if st.button("← Back"):
            st.session_state.step = 2
            st.rerun()
    with col2:
        if st.button("✦ Generate My Tailored CV", use_container_width=True):
            with st.spinner("🤖 Claude is crafting your tailored CV…"):
                try:
                    result = call_claude(
                        st.session_state.cv_text,
                        st.session_state.jd_text,
                        st.session_state.mode,
                    )
                    st.session_state.tailored_cv = result
                    st.session_state.step = 4
                    st.rerun()
                except anthropic.AuthenticationError:
                    st.error("Invalid Anthropic API key. Check your Streamlit Secrets.")
                except Exception as e:
                    st.error(f"Generation failed: {e}")


# ── STEP 4: Result & Payment Gate ─────────────────────────
def page_step4():
    st.markdown("""
    <div class="card">
      <div class="result-header">
        <div class="card-title" style="margin-bottom:0">📋 Your Tailored CV Preview</div>
        <span class="badge">{mode}</span>
      </div>
    </div>
    """.format(mode=st.session_state.mode), unsafe_allow_html=True)

    # Editable text area — user can fine-tune before downloading
    edited = st.text_area(
        "Edit if needed:",
        value=st.session_state.tailored_cv,
        height=400,
        label_visibility="collapsed",
    )
    st.session_state.tailored_cv = edited

    st.markdown('<span class="badge">✓ ATS-Ready</span><span class="badge">✓ Action Verbs</span><span class="badge">✓ Quantified Impact</span>', unsafe_allow_html=True)

    st.markdown("---")

    # ── Payment Gate ──────────────────────────────────────
    if not st.session_state.payment_verified:
        st.markdown("""
        <div class="warn-box">
          🔒 Your CV is ready! Pay <strong>₹20</strong> to unlock the Word (.docx) download.
        </div>
        """, unsafe_allow_html=True)

        razorpay_link = get_razorpay_link()

        st.markdown(f"""
        <div class="pay-btn-wrap">
          <a href="{razorpay_link}" target="_blank" class="pay-btn">
            💳 &nbsp; Pay ₹20 &amp; Unlock Download
          </a>
        </div>
        <p style="text-align:center;font-size:0.78rem;color:#6B7280;margin-top:0.5rem;">
          Secure payment via Razorpay · UPI, Cards, Net Banking accepted
        </p>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("**Already paid?** Click the button below to confirm and unlock your download.")

        col_a, col_b, col_c = st.columns([1, 2, 1])
        with col_b:
            if st.button("✅ I've Paid — Unlock Download", use_container_width=True):
                # In production: verify via Razorpay Webhook / API signature check
                # For MVP: trust-based unlock (replace with real verification)
                st.session_state.payment_verified = True
                st.rerun()

    else:
        # ── Payment Confirmed — Show Download ─────────────
        st.markdown("""
        <div class="success-box">
          ✅ Payment confirmed! Your Word document is ready to download.
        </div>
        """, unsafe_allow_html=True)

        docx_bytes = generate_docx(
            st.session_state.tailored_cv,
            st.session_state.mode,
        )

        st.download_button(
            label="⬇️  Download Tailored CV (.docx)",
            data=docx_bytes,
            file_name=f"TailorMyCV_{st.session_state.mode.replace(' ','-')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.markdown("""
        <div class="info-box">
          💡 <strong>Tip:</strong> Open in Microsoft Word or Google Docs for final formatting tweaks.
          For ATS uploads, save as PDF from Word.
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    col1, col2 = st.columns([1, 2])
    with col1:
        if st.button("← Regenerate"):
            st.session_state.step = 3
            st.session_state.payment_verified = False
            st.session_state.tailored_cv = ""
            st.rerun()
    with col2:
        if st.button("✦ Start Fresh (New CV)", use_container_width=True):
            for k, v in DEFAULTS.items():
                st.session_state[k] = v
            st.rerun()


# ════════════════════════════════════════════════════════════
#  MAIN ROUTER
# ════════════════════════════════════════════════════════════
def main():
    render_masthead()
    render_step_bar(st.session_state.step)

    if st.session_state.step == 1:
        page_step1()
    elif st.session_state.step == 2:
        page_step2()
    elif st.session_state.step == 3:
        page_step3()
    elif st.session_state.step == 4:
        page_step4()

    # Footer
    st.markdown("""
    <div class="footer">
      Made with ❤️ in India · TailorMyCV © 2025 ·
      <a href="mailto:support@tailormycv.in" style="color:#4F46E5;">support@tailormycv.in</a>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
