# ================================================================
#  TailorMyCV V3 — Professionalism & Integrity Edition
#  app_v2.py  (filename as requested)
#  Stack: Streamlit · Claude claude-3-sonnet-20240229 · Razorpay
#         python-docx · reportlab
# ================================================================

import streamlit as st
import anthropic
import io, os, re, random

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import PyPDF2
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    import docx2txt
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph,
                                    Spacer, HRFlowable, KeepTogether)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False


# ════════════════════════════════════════════════════════════
#  PAGE CONFIG
# ════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="TailorMyCV — AI Resume Tailoring",
    page_icon="◆",
    layout="centered",
    initial_sidebar_state="collapsed",
)


# ════════════════════════════════════════════════════════════
#  GLOBAL CSS
# ════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --navy:        #0D1B2A;
  --navy-2:      #132238;
  --navy-3:      #1A2F4A;
  --navy-card:   #1E3550;
  --gold:        #C9A84C;
  --gold-lt:     #E8C97A;
  --gold-dim:    rgba(201,168,76,0.15);
  --teal:        #2DD4BF;
  --teal-dim:    rgba(45,212,191,0.12);
  --text-prime:  #F0EDE8;
  --text-muted:  #8FA3B8;
  --text-dim:    #566D84;
  --border:      rgba(201,168,76,0.18);
  --border-soft: rgba(240,237,232,0.08);
  --danger:      #F87171;
  --success:     #34D399;
  --blue:        #3B82F6;
  --red:         #EF4444;
  --radius:      12px;
  --radius-lg:   18px;
  --shadow:      0 8px 32px rgba(0,0,0,0.40);
  --shadow-gold: 0 4px 20px rgba(201,168,76,0.25);
}

html, body, [class*="css"], .stApp {
  font-family: 'Inter', sans-serif !important;
  background-color: var(--navy) !important;
  color: var(--text-prime) !important;
}
#MainMenu, footer, header { visibility: hidden !important; }
.block-container { padding: 1.5rem 1rem 5rem !important; max-width: 800px !important; }

/* ── MASTHEAD ── */
.masthead {
  text-align: center;
  padding: 2.4rem 1.5rem 1.8rem;
  background: linear-gradient(160deg, var(--navy-3) 0%, var(--navy-2) 100%);
  border-radius: var(--radius-lg);
  border: 1px solid var(--border);
  margin-bottom: 2rem;
  box-shadow: var(--shadow);
  position: relative;
  overflow: hidden;
}
.masthead::before {
  content: '';
  position: absolute; top: -60px; right: -60px;
  width: 200px; height: 200px;
  background: radial-gradient(circle, rgba(201,168,76,0.08) 0%, transparent 70%);
  border-radius: 50%;
}
.masthead-logo {
  font-family: 'Playfair Display', serif;
  font-size: 2.4rem; color: var(--text-prime);
  line-height: 1.1; margin-bottom: 0.45rem;
}
.masthead-logo .accent { color: var(--gold); }
.masthead-sub { font-size: 0.84rem; color: var(--text-muted); letter-spacing: 0.03em; }
.masthead-badges { display:flex; gap:8px; justify-content:center; flex-wrap:wrap; margin-top:1.1rem; }
.mbadge { background:var(--gold-dim); color:var(--gold-lt); border:1px solid rgba(201,168,76,0.3); border-radius:999px; font-size:0.7rem; font-weight:600; padding:4px 12px; }

/* ── STEP BAR ── */
.step-bar { display:flex; align-items:center; justify-content:center; gap:0; margin-bottom:2rem; flex-wrap:wrap; row-gap:6px; }
.step-pill { display:flex; align-items:center; gap:5px; padding:6px 12px; border-radius:999px; font-size:0.7rem; font-weight:600; color:var(--text-dim); background:var(--navy-card); border:1px solid var(--border-soft); transition:all .3s; white-space:nowrap; }
.step-pill.active { background:var(--gold-dim); color:var(--gold-lt); border-color:var(--border); box-shadow:var(--shadow-gold); }
.step-pill.done { background:var(--teal-dim); color:var(--teal); border-color:rgba(45,212,191,0.25); }
.step-connector { width:18px; height:1px; background:var(--border-soft); }

/* ── CARDS ── */
.card { background:var(--navy-card); border-radius:var(--radius-lg); border:1px solid var(--border-soft); padding:1.6rem 1.5rem; margin-bottom:1.1rem; box-shadow:var(--shadow); }
.card-title { font-family:'Playfair Display',serif; font-size:1.05rem; color:var(--gold-lt); margin-bottom:0.85rem; display:flex; align-items:center; gap:8px; }
.divider { height:1px; background:var(--border-soft); margin:1rem 0; }

/* ── STATUS BOXES ── */
.info-box   { background:rgba(45,212,191,0.07); border-left:3px solid var(--teal); border-radius:0 8px 8px 0; padding:9px 13px; font-size:0.81rem; color:var(--teal); margin:0.7rem 0; }
.warn-box   { background:rgba(248,113,113,0.08); border-left:3px solid var(--danger); border-radius:0 8px 8px 0; padding:9px 13px; font-size:0.81rem; color:#FCA5A5; margin:0.7rem 0; }
.success-box{ background:rgba(52,211,153,0.08); border-left:3px solid var(--success); border-radius:0 8px 8px 0; padding:9px 13px; font-size:0.81rem; color:var(--success); margin:0.7rem 0; }
.gold-box   { background:var(--gold-dim); border:1px solid rgba(201,168,76,0.30); border-radius:var(--radius); padding:13px 16px; font-size:0.83rem; color:var(--gold-lt); margin:0.7rem 0; }
.lead-box   { background:rgba(59,130,246,0.10); border:1px solid rgba(59,130,246,0.30); border-radius:var(--radius); padding:13px 16px; font-size:0.82rem; color:#93C5FD; margin:0.7rem 0; }

/* ── TEXT INPUTS — CONTRAST-FIXED ── */
.stTextArea textarea, textarea {
  background-color: #0A1520 !important;
  color: #F0EDE8 !important;
  caret-color: #C9A84C !important;
  border: 1px solid rgba(201,168,76,0.30) !important;
  border-radius: 10px !important;
  font-family: 'JetBrains Mono', monospace !important;
  font-size: 0.82rem !important;
  line-height: 1.55 !important;
  padding: 10px 12px !important;
  box-shadow: inset 0 2px 6px rgba(0,0,0,0.30) !important;
  transition: border-color .2s !important;
  resize: vertical !important;
}
.stTextArea textarea:focus, textarea:focus {
  border-color: #C9A84C !important;
  box-shadow: 0 0 0 3px rgba(201,168,76,0.14), inset 0 2px 6px rgba(0,0,0,0.30) !important;
  outline: none !important;
}
.stTextArea textarea::placeholder, textarea::placeholder { color:#566D84 !important; opacity:1 !important; }
.stTextArea textarea:disabled, textarea:disabled {
  background-color: #060E18 !important;
  color: #3D5568 !important;
  caret-color: transparent !important;
  border-color: rgba(255,255,255,0.04) !important;
  cursor: not-allowed !important;
}
.stTextInput input, input[type="text"], input[type="email"] {
  background-color: #0A1520 !important;
  color: #F0EDE8 !important;
  caret-color: #C9A84C !important;
  border: 1px solid rgba(201,168,76,0.30) !important;
  border-radius: 8px !important;
  font-family: 'Inter', sans-serif !important;
  font-size: 0.87rem !important;
  padding: 9px 13px !important;
}
.stTextInput input:focus, input[type="text"]:focus, input[type="email"]:focus {
  border-color: #C9A84C !important;
  box-shadow: 0 0 0 3px rgba(201,168,76,0.14) !important;
  outline: none !important;
}
.stTextInput input::placeholder { color:#566D84 !important; }
.stTextArea label, .stTextInput label { color:var(--text-muted) !important; font-size:0.81rem !important; font-weight:500 !important; }

/* ── FILE UPLOADER ── */
.stFileUploader section, div[data-testid="stFileUploadDropzone"] {
  background: rgba(201,168,76,0.05) !important;
  border: 2px dashed rgba(201,168,76,0.35) !important;
  border-radius: var(--radius) !important;
}
.stFileUploader section:hover, div[data-testid="stFileUploadDropzone"]:hover {
  background: var(--gold-dim) !important; border-color: var(--gold) !important;
}
.stFileUploader label, .stFileUploader p, .stFileUploader span { color:var(--text-muted) !important; }
[data-testid="stFileUploadDropzone"] p { color:var(--gold-lt) !important; }

/* ── PRIMARY BUTTON (gold) ── */
.stButton > button {
  background: linear-gradient(135deg, #C9A84C 0%, #A8793A 100%) !important;
  color: #0D1B2A !important;
  border: none !important;
  border-radius: 999px !important;
  padding: 0.6rem 1.8rem !important;
  font-family: 'Inter', sans-serif !important;
  font-weight: 700 !important;
  font-size: 0.87rem !important;
  letter-spacing: 0.03em !important;
  transition: all .2s !important;
  box-shadow: 0 4px 18px rgba(201,168,76,0.25) !important;
}
.stButton > button:hover {
  background: linear-gradient(135deg, #E8C97A 0%, #C9A84C 100%) !important;
  transform: translateY(-1px) !important;
  box-shadow: 0 6px 22px rgba(201,168,76,0.38) !important;
}

/* ── TEAL/GREEN DOWNLOAD BUTTON (Word) ── */
div[data-testid="stDownloadButton"]:nth-of-type(1) > button,
.dl-word .stDownloadButton > button {
  background: linear-gradient(135deg, #0D9488 0%, #059669 100%) !important;
  color: #fff !important;
  border: 1px solid rgba(45,212,191,0.40) !important;
  border-radius: 999px !important;
  padding: 0.6rem 1.8rem !important;
  font-family: 'Inter', sans-serif !important;
  font-weight: 700 !important;
  font-size: 0.87rem !important;
  box-shadow: 0 4px 16px rgba(13,148,136,0.30) !important;
  transition: all .2s !important;
}
div[data-testid="stDownloadButton"]:nth-of-type(1) > button:hover,
.dl-word .stDownloadButton > button:hover {
  background: linear-gradient(135deg, #14B8A6 0%, #10B981 100%) !important;
  transform: translateY(-1px) !important;
  box-shadow: 0 8px 22px rgba(13,148,136,0.42) !important;
}

/* ── RED/CORAL DOWNLOAD BUTTON (PDF) ── */
div[data-testid="stDownloadButton"]:nth-of-type(2) > button,
.dl-pdf .stDownloadButton > button {
  background: linear-gradient(135deg, #EF4444 0%, #F97316 100%) !important;
  color: #fff !important;
  border: 1px solid rgba(239,68,68,0.40) !important;
  border-radius: 999px !important;
  padding: 0.6rem 1.8rem !important;
  font-family: 'Inter', sans-serif !important;
  font-weight: 700 !important;
  font-size: 0.87rem !important;
  box-shadow: 0 4px 16px rgba(239,68,68,0.28) !important;
  transition: all .2s !important;
}
div[data-testid="stDownloadButton"]:nth-of-type(2) > button:hover,
.dl-pdf .stDownloadButton > button:hover {
  background: linear-gradient(135deg, #F87171 0%, #FB923C 100%) !important;
  transform: translateY(-1px) !important;
  box-shadow: 0 8px 22px rgba(239,68,68,0.38) !important;
}

/* ── PAY BUTTON ── */
.pay-btn-wrap { text-align:center; margin:1.2rem 0; }
.pay-btn {
  display:inline-block;
  background:linear-gradient(135deg,#C9A84C 0%,#A8793A 100%);
  color:#0D1B2A !important;
  font-family:'Inter',sans-serif;
  font-weight:800; font-size:1rem;
  padding:14px 40px; border-radius:999px;
  text-decoration:none !important;
  box-shadow:0 4px 20px rgba(201,168,76,0.32);
  transition:all .2s; letter-spacing:0.02em;
}
.pay-btn:hover { background:linear-gradient(135deg,#E8C97A 0%,#C9A84C 100%); transform:translateY(-2px); box-shadow:0 8px 26px rgba(201,168,76,0.48); }

/* ── ATS SCORE ── */
.ats-wrap { background:var(--navy-3); border:1px solid var(--border); border-radius:var(--radius-lg); padding:1.3rem 1.5rem; margin:0.9rem 0; }
.ats-header { display:flex; justify-content:space-between; align-items:center; margin-bottom:0.7rem; }
.ats-label { font-size:0.76rem; font-weight:600; letter-spacing:0.12em; text-transform:uppercase; color:var(--text-muted); }
.ats-num { font-family:'Playfair Display',serif; font-size:2.1rem; font-weight:700; line-height:1; }
.ats-num.high { color:var(--success); }
.ats-num.mid  { color:var(--gold); }
.ats-num.low  { color:var(--danger); }
.ats-track { width:100%; height:7px; background:rgba(255,255,255,0.07); border-radius:999px; overflow:hidden; }
.ats-fill  { height:100%; border-radius:999px; }
.ats-gap { margin-top:0.85rem; font-size:0.78rem; color:var(--text-muted); line-height:1.65; }
.ats-gap strong { color:var(--gold-lt); display:block; margin-bottom:0.45rem; }
.ats-gap-list { display:flex; flex-direction:column; gap:5px; margin-top:0.4rem; }
.ats-gap-bullet { background:rgba(248,113,113,0.07); border-left:2px solid rgba(248,113,113,0.45); border-radius:0 6px 6px 0; padding:6px 10px; font-size:0.78rem; color:#FCA5A5; line-height:1.5; }

/* ── BADGES ── */
.badge { display:inline-block; background:var(--gold-dim); color:var(--gold-lt); border:1px solid rgba(201,168,76,0.25); font-size:0.67rem; font-weight:700; letter-spacing:0.05em; padding:3px 9px; border-radius:999px; margin-right:4px; margin-bottom:3px; text-transform:uppercase; }

/* ── MODE CARDS ── */
.mode-card { border:1px solid var(--border-soft); border-radius:var(--radius); padding:13px 9px; text-align:center; background:var(--navy-3); transition:all .2s; }
.mode-card.sel { border-color:var(--gold); background:var(--gold-dim); box-shadow:var(--shadow-gold); }
.mode-icon { font-size:1.45rem; }
.mode-label { font-size:0.76rem; font-weight:700; color:var(--text-prime); margin-top:5px; }
.mode-desc  { font-size:0.67rem; color:var(--text-muted); margin-top:2px; }

/* ── MORALE CARD ── */
.morale-card {
  background:linear-gradient(135deg,var(--navy-3) 0%,#1A3050 100%);
  border:1px solid rgba(201,168,76,0.28);
  border-radius:var(--radius-lg);
  padding:1.7rem; text-align:center;
  margin:1.1rem 0; box-shadow:var(--shadow-gold);
  position:relative; overflow:hidden;
}
.morale-card::after { content:'◆'; position:absolute; bottom:-18px; right:8px; font-size:5.5rem; color:rgba(201,168,76,0.05); line-height:1; }
.morale-icon { font-size:2rem; margin-bottom:0.5rem; }
.morale-text { font-family:'Playfair Display',serif; font-size:1rem; color:var(--gold-lt); line-height:1.55; font-style:italic; }

/* ── REFERRAL ── */
.referral-box { background:var(--teal-dim); border:1px solid rgba(45,212,191,0.18); border-radius:var(--radius); padding:0.95rem 1.1rem; text-align:center; margin-top:0.9rem; font-size:0.82rem; color:var(--teal); }
.referral-box strong { color:#fff; }

/* ── SHARE ROW ── */
.share-row-card {
  background: var(--navy-card);
  border: 1px solid var(--border-soft);
  border-radius: var(--radius);
  padding: 0.85rem 1.2rem;
  margin-top: 1rem;
  display: flex;
  align-items: center;
  gap: 14px;
  flex-wrap: wrap;
}
.share-label {
  font-size: 0.75rem;
  font-weight: 600;
  color: var(--text-muted);
  white-space: nowrap;
  letter-spacing: 0.05em;
  text-transform: uppercase;
}
.share-btns { display: flex; gap: 8px; flex-wrap: wrap; }
.share-btn {
  display: inline-flex; align-items: center;
  gap: 5px; padding: 6px 14px;
  border-radius: 999px; font-size: 0.75rem;
  font-weight: 600; cursor: pointer;
  text-decoration: none !important;
  border: none; font-family: 'Inter', sans-serif;
  transition: all .2s; white-space: nowrap;
}
.share-copy {
  background: #2D3748; color: #CBD5E0 !important;
  border: 1px solid rgba(255,255,255,0.1);
}
.share-copy:hover { background: #3D4F65; color: #fff !important; }
.share-wa {
  background: #25D366; color: #fff !important;
  box-shadow: 0 2px 10px rgba(37,211,102,0.3);
}
.share-wa:hover { background: #20bd5a; box-shadow: 0 4px 14px rgba(37,211,102,0.45); }
.share-ig {
  background: linear-gradient(135deg,#f09433,#e6683c,#dc2743,#cc2366,#bc1888);
  color: #fff !important;
  box-shadow: 0 2px 10px rgba(220,39,67,0.3);
}
.share-ig:hover { opacity: 0.88; box-shadow: 0 4px 14px rgba(220,39,67,0.45); }

/* ── CONFETTI CANVAS ── */
#confetti-canvas {
  position:fixed; top:0; left:0;
  width:100vw; height:100vh;
  pointer-events:none; z-index:9999;
}

/* ── MISC ── */
.stSpinner > div { border-top-color:var(--gold) !important; }
hr { border-color:var(--border-soft) !important; }
p, li { color:var(--text-prime) !important; }
</style>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════
#  CONSTANTS
# ════════════════════════════════════════════════════════════
CLAUDE_MODEL = "claude-haiku-4-5-20251001"  # Confirmed active model (April 2026)

MODE_META = {
    "ATS-Optimizer": {
        "icon": "🎯",
        "desc": "Keyword-match to beat ATS filters",
        "instr": "Prioritise ATS keyword optimisation. Mirror EXACT terminology from the JD.",
    },
    "Executive Tone": {
        "icon": "💼",
        "desc": "C-suite gravitas & boardroom authority",
        "instr": "Rewrite with commanding executive authority. Lead with strategic business impact and P&L ownership.",
    },
    "Career Switch": {
        "icon": "🔄",
        "desc": "Pivot your story to a new industry",
        "instr": "Bridge the career gap. Aggressively reframe transferable skills; the Summary must connect past to future.",
    },
}

# Dynamic morale templates — {name} is substituted at render time
MORALE_TEMPLATES = [
    ("🚀", "Bots beaten, keywords matched, {name}! Your CV is now optimized and ready. Now go show them why you're exactly who they've been waiting for!"),
    ("🏆", "The algorithm has been conquered, {name}! Your story is polished, your value is clear. Walk into that interview like you already have the job."),
    ("⚡", "Mission accomplished, {name}! Your CV is ATS-ready and human-ready. The only thing left to do is show up and shine."),
    ("🎯", "You're in the top tier now, {name}! Keywords locked in, impact quantified. Go make the hiring manager wonder how they ever lived without you."),
    ("💎", "{name}, your CV just levelled up. Precision-crafted, ATS-approved, and ready to open doors. This is your moment — go claim it!"),
    ("🔥", "Recruiters scroll fast, {name} — but your CV just made them stop. You've done the work. Now go own the room."),
    ("🌟", "From application to interview, {name}, your CV is your first impression — and it's a brilliant one. Go make the rest count!"),
    ("🦁", "The hard part is done, {name}! Your CV speaks before you even walk in. Now go show them the person behind the page."),
]

DEFAULTS = {
    "step": 1,
    "cv_text": "",
    "jd_text": "",
    "jd_role": "Professional",          # extracted role title for personalisation
    "mode": "ATS-Optimizer",
    "lead_name": "",
    "lead_email": "",
    "lead_captured": False,
    "tailored_cv": "",
    "ats_score": None,
    "ats_color": "mid",
    "ats_bar_color": "#C9A84C",
    "gap_analysis": "",
    "payment_verified": False,
    "downloaded": False,
    "morale_msg": None,
    "confetti_fired": False,
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ════════════════════════════════════════════════════════════
#  HELPERS — PARSERS
# ════════════════════════════════════════════════════════════
def extract_pdf(b: bytes) -> str:
    if not PDF_OK: return ""
    r = PyPDF2.PdfReader(io.BytesIO(b))
    return "\n".join(p.extract_text() or "" for p in r.pages)

def extract_docx_text(b: bytes) -> str:
    if not DOCX_OK: return ""
    return docx2txt.process(io.BytesIO(b))

def _guess_role(jd_text: str) -> str:
    """Best-effort role extraction from first 300 chars of JD."""
    snippet = jd_text[:300]
    for pattern in [
        r"(?:hiring|looking for|seeking|role[:\s]+|position[:\s]+|title[:\s]+)\s*(?:a\s+|an\s+)?([A-Z][^\n,.(]{3,40})",
        r"^([A-Z][^\n,.(]{3,40})(?:\n|$)",
    ]:
        m = re.search(pattern, snippet, re.IGNORECASE | re.MULTILINE)
        if m:
            role = m.group(1).strip().rstrip(".")
            if 3 < len(role) < 50:
                return role
    return "Professional"


# ════════════════════════════════════════════════════════════
#  ANTHROPIC CLIENT
# ════════════════════════════════════════════════════════════
def get_client():
    try:
        key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception:
        key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        st.error("ANTHROPIC_API_KEY not set. Add it in Streamlit → Settings → Secrets.")
        st.stop()
    return anthropic.Anthropic(api_key=key)


# ════════════════════════════════════════════════════════════
#  SYSTEM PROMPT — INTEGRITY-FIRST
# ════════════════════════════════════════════════════════════
def build_system_prompt(mode: str) -> str:
    instr = MODE_META[mode]["instr"]
    return f"""You are a world-class executive recruiter and ATS specialist with 25 years of experience
placing senior professionals at Fortune 500 companies and leading Indian enterprises.

## MODE
{instr}

## INTEGRITY RULES — NON-NEGOTIABLE
1. DO NOT invent or add any skill, tool, technology, qualification, or achievement
   that does not appear in the original CV. Every line must be truthful to the source.
2. If the candidate lacks a skill required by the JD, DO NOT include it in the CV body.
   Instead, list it ONLY inside the "MISSING FOR SUCCESS" section of the ATS block.
3. Your job is to REFRAME and OPTIMISE existing experience — not to fabricate.

## UNIVERSAL STYLE RULES
4. Use elite action verbs: Spearheaded, Architected, Orchestrated, Transformed, Accelerated,
   Championed, Instituted, Galvanised, Propelled.
5. Every experience bullet: [Action Verb] + [Context] + [Measurable Outcome].
6. Structure MUST follow exactly:

   [CANDIDATE FULL NAME]
   [Phone | Email | LinkedIn | City]

   PROFESSIONAL SUMMARY
   [3–4 sentences]

   CORE COMPETENCIES
   [8–12 pipe-separated skills — truthful to CV only]

   PROFESSIONAL EXPERIENCE
   [Company] | [Title] | [Dates]
   • [bullet]

   EDUCATION
   [Degree | Institution | Year]

   CERTIFICATIONS & AWARDS
   [only if present in original CV]

7. Output ONLY the CV text above — no commentary, no markdown, no preamble.
8. After the CV, on a new line, output EXACTLY:

---ATS_ANALYSIS---
SCORE: [integer 0-100]
MISSING: [2-3 sentences starting with "Your CV is missing..." naming specific skills/keywords from the JD that the candidate does NOT have]
---END_ATS---

Do NOT skip the ATS block."""


# ════════════════════════════════════════════════════════════
#  CLAUDE CALL + ATS PARSING
# ════════════════════════════════════════════════════════════
def call_claude(cv_text, jd_text, mode):
    client = get_client()
    msg = (
        f"=== ORIGINAL CV ===\n{cv_text}\n\n"
        f"=== JOB DESCRIPTION ===\n{jd_text}\n\n"
        f"Tailor strictly using the '{mode}' approach. "
        f"Remember: no invented skills. Append the ---ATS_ANALYSIS--- block."
    )
    resp = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=3000,
        system=build_system_prompt(mode),
        messages=[{"role": "user", "content": msg}],
    )
    full = resp.content[0].text

    score, gap, cv_clean = None, "", full
    m = re.search(
        r"---ATS_ANALYSIS---\s*SCORE:\s*(\d+)\s*MISSING:\s*(.*?)\s*---END_ATS---",
        full, re.DOTALL | re.IGNORECASE,
    )
    if m:
        score    = min(100, max(0, int(m.group(1))))
        gap      = m.group(2).strip()
        cv_clean = full[:m.start()].strip()

    if score is None:
        jd_w  = set(re.findall(r'\b[a-z]{4,}\b', jd_text.lower()))
        cv_w  = set(re.findall(r'\b[a-z]{4,}\b', cv_clean.lower()))
        score = min(93, int((len(jd_w & cv_w) / max(len(jd_w), 1)) * 145))
        gap   = "Your CV is missing a detailed gap breakdown — review the JD keywords manually."

    color     = "high" if score >= 75 else ("mid" if score >= 50 else "low")
    bar_color = "#34D399" if color == "high" else ("#C9A84C" if color == "mid" else "#F87171")
    return {"cv_text": cv_clean, "score": score, "color": color,
            "bar_color": bar_color, "gap": gap}


# ════════════════════════════════════════════════════════════
#  DOCUMENT GENERATORS
# ════════════════════════════════════════════════════════════
def _parse_header(cv_text):
    lines = [l.strip() for l in cv_text.split("\n") if l.strip()]
    name = lines[0] if lines else "Candidate"
    contact = ""
    if len(lines) > 1 and ("|" in lines[1] or "@" in lines[1] or "+" in lines[1]):
        contact = lines[1]
    return name, contact

def _add_hr_docx(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'B8964A')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def generate_docx(cv_text: str) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.80)
        sec.bottom_margin = Inches(0.80)
        sec.left_margin   = Inches(0.95)
        sec.right_margin  = Inches(0.95)

    name, contact = _parse_header(cv_text)

    # ── Name — huge, centred ──
    np = doc.add_paragraph()
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.add_run(name.upper())
    nr.bold = True
    nr.font.size = Pt(20)
    nr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    np.paragraph_format.space_before = Pt(0)
    np.paragraph_format.space_after  = Pt(1)

    # ── Contact — centred ──
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(contact)
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(3)

    _add_hr_docx(doc)

    # ── Body ──
    lines = cv_text.split("\n")
    skip  = sum(1 for l in lines if l.strip() in (name, contact, ""))
    # smarter skip: skip until first non-header, non-contact, non-empty line
    in_header = True
    body_start = 0
    for i, line in enumerate(lines):
        s = line.strip()
        if in_header and s in (name, contact, ""):
            body_start = i + 1
        else:
            in_header = False

    for line in lines[body_start:]:
        s = line.strip()
        if not s:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            continue

        is_heading = s.isupper() and len(s) < 55 and not s.startswith("•")
        if is_heading:
            _add_hr_docx(doc)
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x1A, 0x35, 0x6B)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(1)
        elif s.startswith("•") or s.startswith("-"):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(s.lstrip("•- ").strip())
            r.font.size = Pt(9.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent  = Inches(0.22)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Bold company names and date ranges in experience lines (lines with | or date patterns)
            if "|" in s and re.search(r'\d{4}', s):
                parts = [x.strip() for x in s.split("|")]
                for idx_p, part in enumerate(parts):
                    r = p.add_run(part)
                    r.font.size = Pt(9.5)
                    r.bold = True
                    if idx_p < len(parts) - 1:
                        sep = p.add_run(" | ")
                        sep.font.size = Pt(9.5)
                        sep.bold = False
            else:
                r = p.add_run(s)
                r.font.size = Pt(9.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_pdf(cv_text: str) -> bytes:
    if not REPORTLAB_OK:
        return b""
    buf    = io.BytesIO()
    doc_rl = SimpleDocTemplate(buf, pagesize=A4,
        topMargin=2.4*cm, bottomMargin=2.0*cm,
        leftMargin=2.2*cm, rightMargin=2.2*cm)

    NAVY = colors.HexColor("#1A2340")
    GOLD = colors.HexColor("#B8964A")
    DARK = colors.HexColor("#222222")
    GREY = colors.HexColor("#555555")
    ss   = getSampleStyleSheet()

    from reportlab.lib.enums import TA_JUSTIFY
    name_sty = ParagraphStyle("N", fontName="Helvetica-Bold",
        fontSize=19, textColor=NAVY, alignment=TA_CENTER,
        spaceAfter=3, spaceBefore=0, leading=22)
    cont_sty = ParagraphStyle("C", fontName="Helvetica",
        fontSize=9, textColor=GREY, alignment=TA_CENTER,
        spaceAfter=8, spaceBefore=0, leading=13)
    sec_sty  = ParagraphStyle("S", fontName="Helvetica-Bold",
        fontSize=9.5, textColor=NAVY, spaceBefore=9,
        spaceAfter=2, leading=12)
    body_sty = ParagraphStyle("B", fontName="Helvetica",
        fontSize=9.5, textColor=DARK, spaceAfter=1,
        spaceBefore=0, leading=13, alignment=TA_JUSTIFY)
    bull_sty = ParagraphStyle("Bul", fontName="Helvetica",
        fontSize=9.5, textColor=DARK, leftIndent=12,
        spaceAfter=1, spaceBefore=0, leading=13, alignment=TA_JUSTIFY)
    exp_sty  = ParagraphStyle("Exp", fontName="Helvetica-Bold",
        fontSize=9.5, textColor=DARK, spaceAfter=1,
        spaceBefore=0, leading=13)

    name, contact = _parse_header(cv_text)
    story = []

    # Clear top margin via explicit spacer
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(name.upper(), name_sty))
    if contact:
        story.append(Paragraph(contact, cont_sty))
    story.append(HRFlowable(width="100%", thickness=1,
                             color=GOLD, spaceBefore=2, spaceAfter=8))

    lines = cv_text.split("\n")
    in_header = True
    body_start = 0
    for i, line in enumerate(lines):
        s = line.strip()
        if in_header and s in (name, contact, ""):
            body_start = i + 1
        else:
            in_header = False

    for line in lines[body_start:]:
        s = line.strip()
        if not s:
            story.append(Spacer(1, 2))
            continue
        if s.isupper() and len(s) < 55 and not s.startswith("•"):
            story.append(HRFlowable(width="100%", thickness=0.4,
                                    color=GOLD, spaceBefore=5, spaceAfter=2))
            story.append(Paragraph(s, sec_sty))
        elif s.startswith("•") or s.startswith("-"):
            story.append(Paragraph(f"• {s.lstrip('•- ').strip()}", bull_sty))
        elif "|" in s and re.search(r'\d{4}', s):
            # Experience header line — bold entire line
            story.append(Paragraph(s.replace("&", "&amp;").replace("<", "&lt;"), exp_sty))
        else:
            story.append(Paragraph(s.replace("&", "&amp;").replace("<", "&lt;"), body_sty))

    doc_rl.build(story)
    buf.seek(0)
    return buf.read()


# ════════════════════════════════════════════════════════════
#  RAZORPAY
# ════════════════════════════════════════════════════════════
def razorpay_link():
    try:
        return st.secrets["RAZORPAY_PAYMENT_LINK"]
    except Exception:
        return os.environ.get("RAZORPAY_PAYMENT_LINK", "https://rzp.io/l/tailormycv-demo")


# ════════════════════════════════════════════════════════════
#  UI COMPONENTS
# ════════════════════════════════════════════════════════════
def render_masthead():
    st.markdown("""
    <div class="masthead">
      <div class="masthead-logo">Tailor<span class="accent">My</span>CV</div>
      <div class="masthead-sub">AI-Powered Resume Intelligence for Success-Driven Professionals</div>
      <div class="masthead-badges">
        <span class="mbadge">Claude AI</span>
        <span class="mbadge">ATS-Scored</span>
        <span class="mbadge">100% Truthful</span>
        <span class="mbadge">₹20 / Download</span>
      </div>
    </div>
    """, unsafe_allow_html=True)


def render_step_bar(current):
    steps = ["Upload CV", "Job Description", "Your Details", "Choose Mode", "Results"]
    pills = ""
    for i, label in enumerate(steps, 1):
        cls  = "done" if i < current else ("active" if i == current else "")
        icon = "✓" if i < current else str(i)
        pills += f'<div class="step-pill {cls}">{icon} {label}</div>'
        if i < len(steps):
            pills += '<div class="step-connector"></div>'
    st.markdown(f'<div class="step-bar">{pills}</div>', unsafe_allow_html=True)


def render_ats(score, color, bar_color, gap):
    # Convert gap text into bullet points
    import re as _re
    sentences = [s.strip() for s in _re.split(r'(?<=[.!?])\s+', gap) if s.strip()]
    if len(sentences) > 1:
        bullets_html = "".join(
            f'<div class="ats-gap-bullet">{"⚠️" if i == 0 else "💡"} {s}</div>'
            for i, s in enumerate(sentences)
        )
    else:
        bullets_html = f'<div class="ats-gap-bullet">⚠️ {gap}</div>'

    st.markdown(f"""
    <div class="ats-wrap">
      <div class="ats-header">
        <div class="ats-label">ATS Compatibility Score</div>
        <div class="ats-num {color}">{score}%</div>
      </div>
      <div class="ats-track">
        <div class="ats-fill" style="width:{score}%;background:{bar_color};"></div>
      </div>
      <div class="ats-gap">
        <strong>What's missing for this role?</strong>
        <div class="ats-gap-list">{bullets_html}</div>
      </div>
    </div>
    """, unsafe_allow_html=True)


CONFETTI_JS = """
<canvas id="confetti-canvas"></canvas>
<script>
(function(){
  var canvas = document.getElementById('confetti-canvas');
  var ctx = canvas.getContext('2d');
  canvas.width = window.innerWidth;
  canvas.height = window.innerHeight;
  var pieces = [];
  var colors = ['#C9A84C','#34D399','#3B82F6','#F87171','#E8C97A','#2DD4BF','#ffffff'];
  for(var i=0;i<160;i++){
    pieces.push({
      x: Math.random()*canvas.width,
      y: Math.random()*canvas.height - canvas.height,
      r: Math.random()*6+3,
      d: Math.random()*160+40,
      color: colors[Math.floor(Math.random()*colors.length)],
      tilt: Math.floor(Math.random()*10)-10,
      tiltAngle: 0, tiltAngleInc: Math.random()*0.07+0.05,
      alpha: 1
    });
  }
  var angle = 0, tick = 0;
  function draw(){
    ctx.clearRect(0,0,canvas.width,canvas.height);
    angle += 0.01;
    tick++;
    for(var i=0;i<pieces.length;i++){
      var p = pieces[i];
      p.tiltAngle += p.tiltAngleInc;
      p.y += (Math.cos(angle+p.d)+1+p.r/6)*1.4;
      p.x += Math.sin(angle)*1.5;
      p.tilt = Math.sin(p.tiltAngle)*12;
      if(tick > 200) p.alpha -= 0.008;
      ctx.globalAlpha = Math.max(0, p.alpha);
      ctx.beginPath();
      ctx.lineWidth = p.r/2;
      ctx.strokeStyle = p.color;
      ctx.moveTo(p.x+p.tilt+p.r/4, p.y);
      ctx.lineTo(p.x+p.tilt, p.y+p.tilt+p.r/4);
      ctx.stroke();
      if(p.y > canvas.height+20){ p.y=-10; p.x=Math.random()*canvas.width; p.alpha=1; }
    }
    ctx.globalAlpha = 1;
    if(tick < 280) requestAnimationFrame(draw);
    else ctx.clearRect(0,0,canvas.width,canvas.height);
  }
  draw();
})();
</script>
"""


# ════════════════════════════════════════════════════════════
#  STEPS
# ════════════════════════════════════════════════════════════

# ── Step 1: Upload CV ──────────────────────────────────────
def page_step1():
    st.markdown('<div class="card"><div class="card-title">📄 Upload or Paste Your CV</div>', unsafe_allow_html=True)
    st.markdown('<div class="info-box">Upload a PDF or DOCX, or paste your text below. Either works perfectly.</div>', unsafe_allow_html=True)

    uploaded = st.file_uploader("Upload CV (PDF or DOCX)", type=["pdf","docx"], label_visibility="collapsed")
    if uploaded:
        raw = uploaded.read()
        txt = extract_pdf(raw) if uploaded.name.lower().endswith(".pdf") else extract_docx_text(raw)
        if txt.strip():
            st.session_state.cv_text = txt
            st.success(f"✅ Extracted **{len(txt.split())}** words from `{uploaded.name}`")
        else:
            st.warning("Could not auto-extract — please paste your CV below.")

    cv_in = st.text_area("Or paste your CV text here:",
        value=st.session_state.cv_text, height=200,
        placeholder="Paste your full CV — name, contact, experience, education, skills…")
    st.session_state.cv_text = cv_in
    st.markdown('</div>', unsafe_allow_html=True)

    if st.button("Continue → Paste Job Description", use_container_width=True):
        if len(st.session_state.cv_text.strip()) < 50:
            st.error("Please provide your CV (minimum 50 characters).")
        else:
            st.session_state.step = 2; st.rerun()


# ── Step 2: Job Description ────────────────────────────────
def page_step2():
    st.markdown('<div class="card"><div class="card-title">🎯 Paste the Job Description</div>', unsafe_allow_html=True)
    st.markdown('<div class="info-box">Copy the full JD from LinkedIn, Naukri, or any job portal. More detail = better match.</div>', unsafe_allow_html=True)

    jd_in = st.text_area("Job Description:",
        value=st.session_state.jd_text, height=220,
        placeholder="Paste the complete job description — role summary, responsibilities, required skills…")
    st.session_state.jd_text = jd_in
    st.markdown('</div>', unsafe_allow_html=True)

    c1, c2 = st.columns([1, 2])
    with c1:
        if st.button("← Back", use_container_width=True):
            st.session_state.step = 1; st.rerun()
    with c2:
        if st.button("Continue → Your Details", use_container_width=True):
            if len(st.session_state.jd_text.strip()) < 50:
                st.error("Please paste the job description (minimum 50 characters).")
            else:
                st.session_state.jd_role = _guess_role(st.session_state.jd_text)
                st.session_state.step = 3; st.rerun()


# ── Step 3: Lead Capture ───────────────────────────────────
def page_step3():
    st.markdown('<div class="card"><div class="card-title">✉️ Where should we send your career insights?</div>', unsafe_allow_html=True)
    st.markdown('<div class="info-box">Enter your details below — we\'ll personalise your experience and may send you exclusive career tips.</div>', unsafe_allow_html=True)

    name_in  = st.text_input("Full Name *",
        value=st.session_state.lead_name,
        placeholder="e.g. Priya Sharma")
    email_in = st.text_input("Email Address *",
        value=st.session_state.lead_email,
        placeholder="e.g. priya@gmail.com")

    st.session_state.lead_name  = name_in.strip()
    st.session_state.lead_email = email_in.strip()
    st.markdown('</div>', unsafe_allow_html=True)

    c1, c2 = st.columns([1, 2])
    with c1:
        if st.button("← Back", use_container_width=True):
            st.session_state.step = 2; st.rerun()
    with c2:
        if st.button("Continue → Choose Mode", use_container_width=True):
            if not st.session_state.lead_name:
                st.error("Please enter your full name."); return
            email_pat = r'^[^@\s]+@[^@\s]+\.[^@\s]+$'
            if not re.match(email_pat, st.session_state.lead_email):
                st.error("Please enter a valid email address."); return
            st.session_state.lead_captured = True
            st.session_state.step = 4; st.rerun()

    # Lead confirmation (visible to tool owner in the deployed app logs / UI)
    if st.session_state.lead_captured:
        st.markdown(f"""
        <div class="lead-box">
          ✅ Lead Captured: <strong>{st.session_state.lead_name}</strong>
          — {st.session_state.lead_email}
        </div>
        """, unsafe_allow_html=True)


# ── Step 4: Tailoring Mode ─────────────────────────────────
def page_step4():
    st.markdown('<div class="card"><div class="card-title">⚙️ Select Your Tailoring Strategy</div>', unsafe_allow_html=True)

    cols = st.columns(3)
    for idx, (key, meta) in enumerate(MODE_META.items()):
        with cols[idx]:
            sel = st.session_state.mode == key
            st.markdown(f"""
            <div class="mode-card {'sel' if sel else ''}">
              <div class="mode-icon">{meta['icon']}</div>
              <div class="mode-label">{key}</div>
              <div class="mode-desc">{meta['desc']}</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("✓ Selected" if sel else "Select", key=f"m_{key}", use_container_width=True):
                st.session_state.mode = key; st.rerun()

    st.markdown(f"""
    <div class="gold-box" style="margin-top:0.9rem;">
      ◆ &nbsp;<strong>{st.session_state.mode}</strong> — {MODE_META[st.session_state.mode]['desc']}
    </div>
    """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    c1, c2 = st.columns([1, 2])
    with c1:
        if st.button("← Back", use_container_width=True):
            st.session_state.step = 3; st.rerun()
    with c2:
        if st.button("◆ Generate My Tailored CV", use_container_width=True):
            with st.spinner("🤖 Analysing your CV against the job description…"):
                try:
                    result = call_claude(
                        st.session_state.cv_text,
                        st.session_state.jd_text,
                        st.session_state.mode,
                    )
                    st.session_state.tailored_cv   = result["cv_text"]
                    st.session_state.ats_score     = result["score"]
                    st.session_state.ats_color     = result["color"]
                    st.session_state.ats_bar_color = result["bar_color"]
                    st.session_state.gap_analysis  = result["gap"]
                    st.session_state.step          = 5
                    st.rerun()
                except anthropic.AuthenticationError:
                    st.error("❌ Invalid API key — check Streamlit Secrets → ANTHROPIC_API_KEY.")
                except anthropic.NotFoundError:
                    st.error(f"❌ Model '{CLAUDE_MODEL}' unavailable. Check your Anthropic account.")
                except Exception as e:
                    st.error(f"❌ Generation failed: {e}")


# ── Step 5: Results & Payment Gate ────────────────────────
def page_step5():
    # Confetti on first load
    if not st.session_state.confetti_fired:
        st.markdown(CONFETTI_JS, unsafe_allow_html=True)
        st.session_state.confetti_fired = True

    # ATS Score — always visible
    if st.session_state.ats_score is not None:
        render_ats(
            st.session_state.ats_score,
            st.session_state.ats_color,
            st.session_state.ats_bar_color,
            st.session_state.gap_analysis,
        )

    st.markdown(f"""
    <div class="card" style="margin-bottom:0.5rem;">
      <div class="card-title">
        📋 Your Tailored CV
        <span class="badge">{st.session_state.mode}</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── PAYMENT GATE ──────────────────────────────────────
    if not st.session_state.payment_verified:
        preview = st.session_state.tailored_cv[:300] + "\n\n… [Full content locked — unlock below]"
        st.text_area("Preview (locked):", value=preview, height=130,
                     disabled=True, label_visibility="collapsed")

        st.markdown("""
        <div class="warn-box">
          🔒 Your CV is ready and ATS-scored. Pay <strong>₹20</strong> to unlock
          the full editable content + Word &amp; PDF downloads.
        </div>
        """, unsafe_allow_html=True)

        st.markdown(f"""
        <div class="pay-btn-wrap">
          <a href="{razorpay_link()}" target="_blank" class="pay-btn">
            💳 &nbsp; Pay ₹20 — Unlock Full CV + Downloads
          </a>
        </div>
        <p style="text-align:center;font-size:0.75rem;color:#566D84;margin-top:0.55rem;">
          Secure · Razorpay · UPI · Cards · Net Banking · Wallets
        </p>
        """, unsafe_allow_html=True)

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown("**Already paid?** Confirm below to unlock your files.")
        _, mid, _ = st.columns([1, 2, 1])
        with mid:
            if st.button("✅ I've Paid — Unlock Now", use_container_width=True):
                st.session_state.payment_verified = True; st.rerun()

    else:
        # ════ UNLOCKED ════════════════════════════════════
        st.markdown(f'<div class="success-box">✅ Unlocked, {st.session_state.lead_name}! Edit below, then download.</div>', unsafe_allow_html=True)

        edited = st.text_area("Fine-tune your CV if needed:",
                               value=st.session_state.tailored_cv, height=400)
        st.session_state.tailored_cv = edited

        st.markdown(
            '<span class="badge">✓ ATS-Optimised</span>'
            '<span class="badge">✓ Truthful to Original</span>'
            '<span class="badge">✓ Action Verbs</span>'
            '<span class="badge">✓ Quantified Impact</span>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

        # ATS tip (updated wording)
        st.markdown("""
        <div class="warn-box">
          💡 <strong>Tip:</strong> We recommend using the <strong>Word (.docx)</strong> file
          for application portals for better results. Use PDF for email or direct submissions.
        </div>
        """, unsafe_allow_html=True)

        # ── Download buttons ───────────────────────────────
        fname = st.session_state.mode.replace(" ", "-")
        c_docx, c_pdf = st.columns(2)

        with c_docx:
            st.markdown('<div class="dl-word">', unsafe_allow_html=True)
            docx_bytes = generate_docx(st.session_state.tailored_cv)
            if st.download_button(
                "📄  Download Word (.docx)",
                data=docx_bytes,
                file_name=f"TailorMyCV_{fname}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            ):
                st.session_state.downloaded = True
            st.markdown('</div>', unsafe_allow_html=True)

        with c_pdf:
            st.markdown('<div class="dl-pdf">', unsafe_allow_html=True)
            if REPORTLAB_OK:
                pdf_bytes = generate_pdf(st.session_state.tailored_cv)
                if st.download_button(
                    "📥  Download PDF",
                    data=pdf_bytes,
                    file_name=f"TailorMyCV_{fname}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                ):
                    st.session_state.downloaded = True
            else:
                st.markdown('<div class="info-box">Add <code>reportlab</code> to requirements.txt for PDF export.</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        # ── Morale Booster (role-personalised) ────────────
        if st.session_state.downloaded:
            if st.session_state.morale_msg is None:
                icon, tpl = random.choice(MORALE_TEMPLATES)
                name = st.session_state.lead_name or "Champion"
                st.session_state.morale_msg = (icon, tpl.format(name=name))
            icon, msg = st.session_state.morale_msg
            st.markdown(f"""
            <div class="morale-card">
              <div class="morale-icon">{icon}</div>
              <div class="morale-text">"{msg}"</div>
            </div>
            """, unsafe_allow_html=True)

            # Sharing row
            APP_URL = "https://tailormycv.streamlit.app"
            wa_msg  = f"I just optimised my CV with TailorMyCV — AI-powered, ATS-scored, ₹20 only! Try it: {APP_URL}"
            wa_url  = f"https://wa.me/?text={wa_msg.replace(' ', '%20')}"
            ig_url  = f"https://www.instagram.com/"
            st.markdown(f"""
            <div class="share-row-card">
              <span class="share-label">Share TailorMyCV</span>
              <div class="share-btns">
                <button class="share-btn share-copy" onclick="navigator.clipboard.writeText('{APP_URL}');this.innerText='✅ Copied!';setTimeout(()=>this.innerHTML='🔗 Copy Link',1800)">🔗 Copy Link</button>
                <a href="{wa_url}" target="_blank" class="share-btn share-wa">
                  <svg width="15" height="15" viewBox="0 0 32 32" fill="white" xmlns="http://www.w3.org/2000/svg" style="vertical-align:middle;margin-right:4px"><path d="M16 2C8.268 2 2 8.268 2 16c0 2.478.677 4.8 1.855 6.793L2 30l7.393-1.836A13.93 13.93 0 0 0 16 30c7.732 0 14-6.268 14-14S23.732 2 16 2zm0 25.6a11.55 11.55 0 0 1-5.887-1.608l-.422-.25-4.387 1.09 1.107-4.275-.275-.44A11.558 11.558 0 0 1 4.4 16C4.4 9.593 9.593 4.4 16 4.4c6.407 0 11.6 5.193 11.6 11.6 0 6.407-5.193 11.6-11.6 11.6zm6.354-8.687c-.348-.174-2.06-1.016-2.38-1.132-.32-.116-.553-.174-.786.174-.232.348-.9 1.132-1.104 1.365-.203.232-.406.26-.754.086-.348-.174-1.47-.542-2.8-1.727-1.034-.924-1.732-2.065-1.935-2.413-.203-.348-.022-.536.153-.71.156-.155.348-.406.522-.61.174-.203.232-.348.348-.58.116-.232.058-.435-.029-.61-.087-.174-.786-1.897-1.078-2.598-.283-.681-.572-.588-.786-.599l-.668-.011c-.232 0-.61.087-.929.435-.319.348-1.22 1.19-1.22 2.903 0 1.714 1.249 3.369 1.423 3.601.174.232 2.458 3.753 5.955 5.263.833.36 1.483.575 1.99.737.836.267 1.598.229 2.2.139.671-.1 2.06-.843 2.35-1.657.29-.813.29-1.51.203-1.657-.086-.145-.319-.232-.667-.406z"/></svg>WhatsApp
                </a>
                <a href="{ig_url}" target="_blank" class="share-btn share-ig">
                  <svg width="15" height="15" viewBox="0 0 32 32" fill="white" xmlns="http://www.w3.org/2000/svg" style="vertical-align:middle;margin-right:4px"><path d="M16 2.88c4.275 0 4.782.017 6.464.093 1.56.071 2.407.33 2.971.548a4.952 4.952 0 0 1 1.838 1.196 4.952 4.952 0 0 1 1.196 1.838c.218.564.477 1.412.548 2.971.076 1.682.093 2.189.093 6.464s-.017 4.782-.093 6.464c-.071 1.56-.33 2.407-.548 2.971a4.952 4.952 0 0 1-1.196 1.838 4.952 4.952 0 0 1-1.838 1.196c-.564.218-1.412.477-2.971.548-1.682.076-2.189.093-6.464.093s-4.782-.017-6.464-.093c-1.56-.071-2.407-.33-2.971-.548a4.952 4.952 0 0 1-1.838-1.196 4.952 4.952 0 0 1-1.196-1.838c-.218-.564-.477-1.412-.548-2.971C2.897 20.782 2.88 20.275 2.88 16s.017-4.782.093-6.464c.071-1.56.33-2.407.548-2.971A4.952 4.952 0 0 1 4.717 4.72 4.952 4.952 0 0 1 6.555 3.52c.564-.218 1.412-.477 2.971-.548C11.218 2.897 11.725 2.88 16 2.88M16 0c-4.346 0-4.89.018-6.596.096-1.702.078-2.863.347-3.88.741a7.833 7.833 0 0 0-2.833 1.843A7.833 7.833 0 0 0 .848 5.513C.454 6.53.185 7.692.107 9.394.029 11.1.011 11.644.011 15.989s.018 4.89.096 6.596c.078 1.702.347 2.863.741 3.88a7.833 7.833 0 0 0 1.843 2.833 7.833 7.833 0 0 0 2.833 1.843c1.017.394 2.178.663 3.88.741C11.11 31.96 11.654 31.978 16 31.978s4.89-.018 6.596-.096c1.702-.078 2.863-.347 3.88-.741a7.833 7.833 0 0 0 2.833-1.843 7.833 7.833 0 0 0 1.843-2.833c.394-1.017.663-2.178.741-3.88.078-1.706.096-2.25.096-6.596s-.018-4.89-.096-6.596c-.078-1.702-.347-2.863-.741-3.88a7.833 7.833 0 0 0-1.843-2.833A7.833 7.833 0 0 0 26.476.848C25.46.454 24.298.185 22.596.107 20.89.029 20.346.011 16 .011zm0 7.784a8.216 8.216 0 1 0 0 16.432 8.216 8.216 0 0 0 0-16.432zm0 13.549a5.333 5.333 0 1 1 0-10.666 5.333 5.333 0 0 1 0 10.666zm8.538-13.876a1.92 1.92 0 1 1-3.84 0 1.92 1.92 0 0 1 3.84 0z"/></svg>Instagram
                </a>
              </div>
            </div>
            """, unsafe_allow_html=True)

    # ── Navigation ─────────────────────────────────────────
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    c1, c2 = st.columns([1, 2])
    with c1:
        if st.button("← Regenerate", use_container_width=True):
            for k in ("tailored_cv","payment_verified","downloaded",
                      "morale_msg","confetti_fired","ats_score"):
                st.session_state[k] = DEFAULTS[k]
            st.session_state.step = 4; st.rerun()
    with c2:
        if st.button("◆ Start Fresh (New CV)", use_container_width=True):
            for k, v in DEFAULTS.items():
                st.session_state[k] = v
            st.rerun()


# ════════════════════════════════════════════════════════════
#  MAIN ROUTER
# ════════════════════════════════════════════════════════════
def main():
    render_masthead()
    render_step_bar(st.session_state.step)

    if   st.session_state.step == 1: page_step1()
    elif st.session_state.step == 2: page_step2()
    elif st.session_state.step == 3: page_step3()
    elif st.session_state.step == 4: page_step4()
    elif st.session_state.step == 5: page_step5()

if __name__ == "__main__":
    main()
