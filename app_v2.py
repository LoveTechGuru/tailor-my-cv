# ================================================================
#  TailorMyCV V2 — Executive Edition  |  app.py
#  Premium AI Resume Tailoring for Senior Candidates
#  Stack: Streamlit · Claude claude-sonnet-4-5 · Razorpay · python-docx
# ================================================================

import streamlit as st
import anthropic
import io
import os
import re
import random

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Optional parsers ─────────────────────────────────────────
try:
    import PyPDF2
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    import docx2txt
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False


# ════════════════════════════════════════════════════════════
#  PAGE CONFIG  (must be first Streamlit call)
# ════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="TailorMyCV — Executive Edition",
    page_icon="◆",
    layout="centered",
    initial_sidebar_state="collapsed",
)


# ════════════════════════════════════════════════════════════
#  GLOBAL CSS — Dark Navy Executive Theme
# ════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=Inter:wght@300;400;500;600&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --navy:        #0D1B2A;
  --navy-2:      #132238;
  --navy-3:      #1A2F4A;
  --navy-card:   #1E3550;
  --gold:        #C9A84C;
  --gold-lt:     #E8C97A;
  --gold-dim:    rgba(201,168,76,0.15);
  --teal:        #2DD4BF;
  --teal-dim:    rgba(45,212,191,0.12);
  --text-prime:  #F0EDE8;
  --text-muted:  #8FA3B8;
  --text-dim:    #566D84;
  --border:      rgba(201,168,76,0.18);
  --border-soft: rgba(240,237,232,0.08);
  --danger:      #F87171;
  --success:     #34D399;
  --radius:      12px;
  --radius-lg:   18px;
  --shadow:      0 8px 32px rgba(0,0,0,0.40);
  --shadow-gold: 0 4px 20px rgba(201,168,76,0.25);
}

html, body, [class*="css"], .stApp {
  font-family: 'Inter', sans-serif !important;
  background-color: var(--navy) !important;
  color: var(--text-prime) !important;
}

#MainMenu, footer, header { visibility: hidden !important; }
.block-container { padding: 1.5rem 1rem 5rem !important; max-width: 800px !important; }

/* MASTHEAD */
.masthead {
  text-align: center;
  padding: 3rem 1.5rem 2rem;
  background: linear-gradient(160deg, var(--navy-3) 0%, var(--navy-2) 100%);
  border-radius: var(--radius-lg);
  border: 1px solid var(--border);
  margin-bottom: 2rem;
  box-shadow: var(--shadow);
  position: relative;
  overflow: hidden;
}
.masthead::before {
  content: '';
  position: absolute;
  top: -60px; right: -60px;
  width: 200px; height: 200px;
  background: radial-gradient(circle, rgba(201,168,76,0.08) 0%, transparent 70%);
  border-radius: 50%;
}
.masthead-eyebrow { font-size:0.72rem; font-weight:600; letter-spacing:0.22em; text-transform:uppercase; color:var(--gold); margin-bottom:0.6rem; }
.masthead-logo { font-family:'Playfair Display',serif; font-size:2.4rem; color:var(--text-prime); line-height:1.1; margin-bottom:0.5rem; }
.masthead-logo .accent { color:var(--gold); }
.masthead-sub { font-size:0.85rem; color:var(--text-muted); letter-spacing:0.03em; }
.masthead-badges { display:flex; gap:8px; justify-content:center; flex-wrap:wrap; margin-top:1.2rem; }
.mbadge { background:var(--gold-dim); color:var(--gold-lt); border:1px solid rgba(201,168,76,0.3); border-radius:999px; font-size:0.7rem; font-weight:600; padding:4px 12px; letter-spacing:0.04em; }

/* STEP BAR */
.step-bar { display:flex; align-items:center; justify-content:center; gap:0; margin-bottom:2rem; }
.step-pill { display:flex; align-items:center; gap:6px; padding:7px 14px; border-radius:999px; font-size:0.73rem; font-weight:600; color:var(--text-dim); background:var(--navy-card); border:1px solid var(--border-soft); transition:all .3s; white-space:nowrap; }
.step-pill.active { background:var(--gold-dim); color:var(--gold-lt); border-color:var(--border); box-shadow:var(--shadow-gold); }
.step-pill.done { background:var(--teal-dim); color:var(--teal); border-color:rgba(45,212,191,0.25); }
.step-connector { width:22px; height:1px; background:var(--border-soft); }

/* CARDS */
.card { background:var(--navy-card); border-radius:var(--radius-lg); border:1px solid var(--border-soft); padding:1.8rem 1.6rem; margin-bottom:1.2rem; box-shadow:var(--shadow); }
.card-title { font-family:'Playfair Display',serif; font-size:1.1rem; color:var(--gold-lt); margin-bottom:0.9rem; display:flex; align-items:center; gap:8px; }
.divider { height:1px; background:var(--border-soft); margin:1.2rem 0; }

/* INFO BOXES */
.info-box { background:rgba(45,212,191,0.07); border-left:3px solid var(--teal); border-radius:0 8px 8px 0; padding:10px 14px; font-size:0.82rem; color:var(--teal); margin:0.8rem 0; }
.warn-box { background:rgba(248,113,113,0.08); border-left:3px solid var(--danger); border-radius:0 8px 8px 0; padding:10px 14px; font-size:0.82rem; color:#FCA5A5; margin:0.8rem 0; }
.success-box { background:rgba(52,211,153,0.08); border-left:3px solid var(--success); border-radius:0 8px 8px 0; padding:10px 14px; font-size:0.82rem; color:var(--success); margin:0.8rem 0; }
.gold-box { background:var(--gold-dim); border:1px solid rgba(201,168,76,0.30); border-radius:var(--radius); padding:14px 18px; font-size:0.85rem; color:var(--gold-lt); margin:0.8rem 0; }

/* ===== CRITICAL TEXT INPUT CONTRAST FIX ===== */
/* Dark background, light text, gold cursor — on ALL textarea and input elements */
.stTextArea textarea,
div[data-testid="stTextArea"] textarea,
textarea {
  background-color: #0A1520 !important;
  color: #F0EDE8 !important;
  caret-color: #C9A84C !important;
  border: 1px solid rgba(201,168,76,0.30) !important;
  border-radius: 10px !important;
  font-family: 'JetBrains Mono', monospace !important;
  font-size: 0.82rem !important;
  line-height: 1.6 !important;
  padding: 12px 14px !important;
  box-shadow: inset 0 2px 8px rgba(0,0,0,0.35) !important;
  transition: border-color .2s !important;
}
.stTextArea textarea:focus,
textarea:focus {
  border-color: #C9A84C !important;
  box-shadow: 0 0 0 3px rgba(201,168,76,0.15), inset 0 2px 8px rgba(0,0,0,0.35) !important;
  outline: none !important;
}
.stTextArea textarea::placeholder,
textarea::placeholder {
  color: #566D84 !important;
  opacity: 1 !important;
}
/* disabled textarea (locked preview) */
.stTextArea textarea:disabled,
textarea:disabled {
  background-color: #060E18 !important;
  color: #4A6070 !important;
  caret-color: transparent !important;
  border-color: rgba(255,255,255,0.05) !important;
  cursor: not-allowed !important;
}

/* Text input (single line) */
.stTextInput input,
input[type="text"] {
  background-color: #0A1520 !important;
  color: #F0EDE8 !important;
  caret-color: #C9A84C !important;
  border: 1px solid rgba(201,168,76,0.30) !important;
  border-radius: 8px !important;
  font-family: 'Inter', sans-serif !important;
  font-size: 0.88rem !important;
  padding: 10px 14px !important;
}
.stTextInput input:focus { border-color: #C9A84C !important; box-shadow: 0 0 0 3px rgba(201,168,76,0.15) !important; outline: none !important; }
.stTextInput input::placeholder { color: #566D84 !important; }

/* Labels */
.stTextArea label, .stTextInput label,
div[data-testid="stTextArea"] label,
div[data-testid="stTextInput"] label {
  color: var(--text-muted) !important;
  font-size: 0.82rem !important;
  font-weight: 500 !important;
}

/* FILE UPLOADER */
.stFileUploader section,
div[data-testid="stFileUploadDropzone"] {
  background: rgba(201,168,76,0.05) !important;
  border: 2px dashed rgba(201,168,76,0.35) !important;
  border-radius: var(--radius) !important;
}
.stFileUploader section:hover, div[data-testid="stFileUploadDropzone"]:hover {
  background: var(--gold-dim) !important; border-color: var(--gold) !important;
}
.stFileUploader label, .stFileUploader p, .stFileUploader span { color: var(--text-muted) !important; }
[data-testid="stFileUploadDropzone"] p { color: var(--gold-lt) !important; }

/* BUTTONS */
.stButton > button {
  background: linear-gradient(135deg, #C9A84C 0%, #A8793A 100%) !important;
  color: #0D1B2A !important;
  border: none !important;
  border-radius: 999px !important;
  padding: 0.65rem 2rem !important;
  font-family: 'Inter', sans-serif !important;
  font-weight: 700 !important;
  font-size: 0.88rem !important;
  letter-spacing: 0.03em !important;
  transition: all .2s !important;
  box-shadow: 0 4px 20px rgba(201,168,76,0.25) !important;
}
.stButton > button:hover {
  background: linear-gradient(135deg, #E8C97A 0%, #C9A84C 100%) !important;
  transform: translateY(-1px) !important;
  box-shadow: 0 6px 24px rgba(201,168,76,0.40) !important;
}

/* Download buttons */
.stDownloadButton > button {
  background: linear-gradient(135deg, #1A5C4A 0%, #0E3D31 100%) !important;
  color: #34D399 !important;
  border: 1px solid rgba(52,211,153,0.30) !important;
  border-radius: 999px !important;
  padding: 0.65rem 2rem !important;
  font-family: 'Inter', sans-serif !important;
  font-weight: 700 !important;
  font-size: 0.88rem !important;
  box-shadow: 0 4px 18px rgba(52,211,153,0.15) !important;
  transition: all .2s !important;
}
.stDownloadButton > button:hover {
  background: linear-gradient(135deg, #216B57 0%, #144D3E 100%) !important;
  transform: translateY(-1px) !important;
  box-shadow: 0 8px 24px rgba(52,211,153,0.25) !important;
}

/* PAY BUTTON */
.pay-btn-wrap { text-align:center; margin:1.2rem 0; }
.pay-btn {
  display:inline-block;
  background:linear-gradient(135deg,#C9A84C 0%,#A8793A 100%);
  color:#0D1B2A !important;
  font-family:'Inter',sans-serif;
  font-weight:800;
  font-size:1.05rem;
  padding:15px 44px;
  border-radius:999px;
  text-decoration:none !important;
  box-shadow:0 4px 20px rgba(201,168,76,0.35);
  transition:all .2s;
  letter-spacing:0.02em;
}
.pay-btn:hover { background:linear-gradient(135deg,#E8C97A 0%,#C9A84C 100%); transform:translateY(-2px); box-shadow:0 8px 28px rgba(201,168,76,0.50); }

/* ATS SCORE */
.ats-score-wrap { background:var(--navy-3); border:1px solid var(--border); border-radius:var(--radius-lg); padding:1.4rem 1.6rem; margin:1rem 0; }
.ats-score-header { display:flex; justify-content:space-between; align-items:center; margin-bottom:0.8rem; }
.ats-label { font-size:0.78rem; font-weight:600; letter-spacing:0.12em; text-transform:uppercase; color:var(--text-muted); }
.ats-number { font-family:'Playfair Display',serif; font-size:2.2rem; font-weight:700; line-height:1; }
.ats-number.high { color:var(--success); }
.ats-number.mid  { color:var(--gold); }
.ats-number.low  { color:var(--danger); }
.ats-bar-track { width:100%; height:8px; background:rgba(255,255,255,0.07); border-radius:999px; overflow:hidden; }
.ats-bar-fill  { height:100%; border-radius:999px; }
.ats-gap { margin-top:0.9rem; font-size:0.78rem; color:var(--text-muted); line-height:1.6; }
.ats-gap strong { color:var(--gold-lt); }

/* BADGES */
.badge { display:inline-block; background:var(--gold-dim); color:var(--gold-lt); border:1px solid rgba(201,168,76,0.25); font-size:0.68rem; font-weight:700; letter-spacing:0.05em; padding:3px 10px; border-radius:999px; margin-right:5px; margin-bottom:4px; text-transform:uppercase; }

/* MODE CARDS */
.mode-card { border:1px solid var(--border-soft); border-radius:var(--radius); padding:14px 10px; text-align:center; background:var(--navy-3); transition:all .2s; }
.mode-card.selected { border-color:var(--gold); background:var(--gold-dim); box-shadow:var(--shadow-gold); }
.mode-icon { font-size:1.5rem; }
.mode-label { font-size:0.78rem; font-weight:700; color:var(--text-prime); margin-top:6px; }
.mode-desc  { font-size:0.68rem; color:var(--text-muted); margin-top:3px; }

/* MORALE CARD */
.morale-card {
  background:linear-gradient(135deg,var(--navy-3) 0%,#1A3050 100%);
  border:1px solid rgba(201,168,76,0.30);
  border-radius:var(--radius-lg);
  padding:1.8rem;
  text-align:center;
  margin:1.2rem 0;
  box-shadow:var(--shadow-gold);
  position:relative;
  overflow:hidden;
}
.morale-card::after { content:'◆'; position:absolute; bottom:-20px; right:10px; font-size:6rem; color:rgba(201,168,76,0.05); line-height:1; }
.morale-icon { font-size:2.2rem; margin-bottom:0.6rem; }
.morale-text { font-family:'Playfair Display',serif; font-size:1.05rem; color:var(--gold-lt); line-height:1.5; font-style:italic; }

/* REFERRAL */
.referral-box { background:var(--teal-dim); border:1px solid rgba(45,212,191,0.20); border-radius:var(--radius); padding:1rem 1.2rem; text-align:center; margin-top:1rem; font-size:0.83rem; color:var(--teal); }
.referral-box strong { color:#fff; }

/* FOOTER */
.footer { text-align:center; font-size:0.73rem; color:var(--text-dim); margin-top:3rem; padding-top:1.2rem; border-top:1px solid var(--border-soft); }
.footer a { color:var(--gold); text-decoration:none; }

/* Misc */
.stSpinner > div { border-top-color:var(--gold) !important; }
hr { border-color:var(--border-soft) !important; }
p, li { color:var(--text-prime) !important; }
</style>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════
#  CONSTANTS & SESSION STATE
# ════════════════════════════════════════════════════════════
MORALE_MESSAGES = [
    ("🚀", "You've got the skills — and now the ATS knows it too. Go crush that interview!"),
    ("🏆", "Top 1% of applicants are made, not born. You just made the cut. Own the room."),
    ("⚡", "Every hiring manager scanning this document is about to see exactly why you're the one."),
    ("🎯", "Precision-engineered for impact. Your story is now impossible to ignore."),
    ("💎", "Diamonds don't beg for attention. Neither does this CV. Walk in with confidence."),
    ("🔥", "The bots have been beaten. Now go show the humans what you're made of."),
    ("🌟", "This isn't just a resume — it's a first-class ticket to the interview room. Board up!"),
    ("🦁", "Senior talent deserves senior presentation. You now have both. The role is yours to lose."),
]

DEFAULTS = {
    "step": 1,
    "cv_text": "",
    "jd_text": "",
    "mode": "ATS-Optimizer",
    "tailored_cv": "",
    "ats_score": None,
    "ats_color": "mid",
    "ats_bar_color": "#C9A84C",
    "gap_analysis": "",
    "payment_verified": False,
    "downloaded": False,
    "morale_msg": None,
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ════════════════════════════════════════════════════════════
#  FILE PARSERS
# ════════════════════════════════════════════════════════════
def extract_pdf(file_bytes: bytes) -> str:
    if not PDF_OK:
        return ""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    return "\n".join(p.extract_text() or "" for p in reader.pages)

def extract_docx(file_bytes: bytes) -> str:
    if not DOCX_OK:
        return ""
    return docx2txt.process(io.BytesIO(file_bytes))


# ════════════════════════════════════════════════════════════
#  ANTHROPIC CLIENT
# ════════════════════════════════════════════════════════════
def get_client():
    try:
        key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception:
        key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        st.error("ANTHROPIC_API_KEY not found. Add it in Streamlit → Settings → Secrets.")
        st.stop()
    return anthropic.Anthropic(api_key=key)

CLAUDE_MODEL = "claude-sonnet-4-5"  # Stable Sonnet release

MODE_META = {
    "ATS-Optimizer": {
        "icon": "🎯",
        "desc": "Beat ATS filters with keyword matching",
        "instr": "PRIORITY: ATS keyword optimization. Mirror EXACT terminology, job titles, and technical skills from the JD. Use standard section headings. Every bullet must contain a measurable outcome.",
    },
    "Executive Tone": {
        "icon": "💼",
        "desc": "C-suite gravitas & boardroom authority",
        "instr": "Rewrite with commanding executive authority. Lead every bullet with strategic business impact, P&L ownership, and leadership scope. Eliminate junior-sounding language.",
    },
    "Career Switch": {
        "icon": "🔄",
        "desc": "Pivot your narrative to a new industry",
        "instr": "Write a compelling bridge narrative. Reframe every transferable skill toward the target role. The Summary must explicitly connect past to future. Downplay irrelevant experience.",
    },
}


# ════════════════════════════════════════════════════════════
#  SYSTEM PROMPT
# ════════════════════════════════════════════════════════════
def build_system_prompt(mode: str) -> str:
    return f"""You are a world-class executive recruiter and ATS specialist with 25 years of experience
placing C-suite leaders at Fortune 500 companies and top Indian conglomerates.

MODE DIRECTIVE: {MODE_META[mode]['instr']}

UNIVERSAL RULES:
1. Never fabricate experience — only reframe what exists.
2. Use elite action verbs: Spearheaded, Architected, Orchestrated, Transformed, Accelerated.
3. Document structure:
   [CANDIDATE NAME]
   [Phone | Email | LinkedIn | City]

   PROFESSIONAL SUMMARY
   [3-4 sentence executive summary]

   CORE COMPETENCIES
   [8-12 pipe-separated skills]

   PROFESSIONAL EXPERIENCE
   [Company] | [Title] | [Dates]
   • [Action] + [Context] + [Measurable Result]

   EDUCATION
   [Degree | Institution | Year]

   CERTIFICATIONS & AWARDS (if present)

4. Output ONLY the CV text — no commentary, no markdown headers, no preamble.
5. At the very end, on a new line, append EXACTLY:

---ATS_ANALYSIS---
SCORE: [integer 0-100]
GAPS: [2-3 sentences naming specific missing keywords or skills from the JD]
---END_ATS---

Do NOT omit the ATS block. It is parsed programmatically."""


# ════════════════════════════════════════════════════════════
#  CLAUDE API CALL + ATS PARSING
# ════════════════════════════════════════════════════════════
def call_claude(cv_text: str, jd_text: str, mode: str) -> dict:
    client = get_client()
    user_msg = (
        f"=== ORIGINAL CV ===\n{cv_text}\n\n"
        f"=== JOB DESCRIPTION ===\n{jd_text}\n\n"
        f"Tailor using the '{mode}' approach. Append the ---ATS_ANALYSIS--- block."
    )
    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=3000,
        system=build_system_prompt(mode),
        messages=[{"role": "user", "content": user_msg}],
    )
    full_text = response.content[0].text

    score, gap, cv_clean = None, "", full_text
    match = re.search(
        r"---ATS_ANALYSIS---\s*SCORE:\s*(\d+)\s*GAPS:\s*(.*?)\s*---END_ATS---",
        full_text, re.DOTALL | re.IGNORECASE,
    )
    if match:
        score    = min(100, max(0, int(match.group(1))))
        gap      = match.group(2).strip()
        cv_clean = full_text[:match.start()].strip()

    if score is None:
        jd_w  = set(re.findall(r'\b[a-z]{4,}\b', jd_text.lower()))
        cv_w  = set(re.findall(r'\b[a-z]{4,}\b', cv_clean.lower()))
        score = min(95, int((len(jd_w & cv_w) / max(len(jd_w), 1)) * 150))
        gap   = "Full gap analysis unavailable — review JD keywords manually."

    color     = "high" if score >= 75 else ("mid" if score >= 50 else "low")
    bar_color = "#34D399" if color == "high" else ("#C9A84C" if color == "mid" else "#F87171")
    return {"cv_text": cv_clean, "score": score, "color": color, "bar_color": bar_color, "gap": gap}


# ════════════════════════════════════════════════════════════
#  DOCUMENT HELPERS
# ════════════════════════════════════════════════════════════
def _parse_name_contact(cv_text: str):
    lines = [l.strip() for l in cv_text.split("\n") if l.strip()]
    name    = lines[0] if lines else "Candidate"
    contact = ""
    if len(lines) > 1:
        second = lines[1]
        # Contact line typically contains | or @ symbols
        if "|" in second or "@" in second or "+" in second:
            contact = second
    return name, contact

def _add_hr(doc):
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after  = Pt(3)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'B8964A')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def generate_docx(cv_text: str, mode: str) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.85)
        sec.bottom_margin = Inches(0.85)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    name, contact = _parse_name_contact(cv_text)

    # Name — huge, centered
    np = doc.add_paragraph()
    np.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = np.add_run(name.upper())
    nr.bold = True
    nr.font.size = Pt(22)
    nr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    np.paragraph_format.space_after = Pt(2)

    # Contact — centered below name
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(contact)
        cr.font.size = Pt(9.5)
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        cp.paragraph_format.space_after = Pt(4)

    _add_hr(doc)

    # Skip name/contact lines in body
    lines = cv_text.split("\n")
    skip = 0
    for line in lines:
        if line.strip() in (name, contact, ""):
            skip += 1
        else:
            break

    for line in lines[skip:]:
        s = line.strip()
        if not s:
            doc.add_paragraph("")
            continue
        is_heading = s.isupper() and len(s) < 50 and not s.startswith("•")
        if is_heading:
            _add_hr(doc)
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(0x1A, 0x35, 0x6B)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
        elif s.startswith("•") or s.startswith("-"):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(s.lstrip("•- ").strip())
            r.font.size = Pt(10)
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_after  = Pt(1)
        else:
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(1)

    # Branding
    doc.add_paragraph("")
    bp = doc.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run(f"◆ TailorMyCV Executive Edition — {mode}")
    br.font.size = Pt(7.5)
    br.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_pdf(cv_text: str, mode: str) -> bytes:
    if not REPORTLAB_OK:
        return b""
    buf    = io.BytesIO()
    doc_rl = SimpleDocTemplate(buf, pagesize=A4,
        topMargin=2.2*cm, bottomMargin=2*cm,
        leftMargin=2.2*cm, rightMargin=2.2*cm)

    NAVY = colors.HexColor("#1A2340")
    GOLD = colors.HexColor("#B8964A")
    DARK = colors.HexColor("#1E1E1E")
    GREY = colors.HexColor("#555555")

    styles = getSampleStyleSheet()
    name_sty = ParagraphStyle("N", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=20, textColor=NAVY,
        alignment=TA_CENTER, spaceAfter=4)
    cont_sty = ParagraphStyle("C", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9, textColor=GREY,
        alignment=TA_CENTER, spaceAfter=6)
    sec_sty  = ParagraphStyle("S", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=10, textColor=NAVY,
        spaceBefore=10, spaceAfter=2, leading=13)
    body_sty = ParagraphStyle("B", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9.5, textColor=DARK,
        spaceAfter=2, leading=14)
    bull_sty = ParagraphStyle("Bul", parent=body_sty,
        leftIndent=14, bulletIndent=4, spaceAfter=1)
    brand_sty = ParagraphStyle("Br", parent=styles["Normal"],
        fontName="Helvetica-Oblique", fontSize=7, textColor=GREY,
        alignment=TA_CENTER, spaceBefore=14)

    name, contact = _parse_name_contact(cv_text)
    story = [Paragraph(name.upper(), name_sty)]
    if contact:
        story.append(Paragraph(contact, cont_sty))
    story.append(HRFlowable(width="100%", thickness=1, color=GOLD, spaceAfter=6))

    lines = cv_text.split("\n")
    skip = 0
    for line in lines:
        if line.strip() in (name, contact, ""):
            skip += 1
        else:
            break

    for line in lines[skip:]:
        s = line.strip()
        if not s:
            story.append(Spacer(1, 4))
            continue
        if s.isupper() and len(s) < 50 and not s.startswith("•"):
            story.append(HRFlowable(width="100%", thickness=0.5, color=GOLD, spaceBefore=6, spaceAfter=3))
            story.append(Paragraph(s, sec_sty))
        elif s.startswith("•") or s.startswith("-"):
            story.append(Paragraph(f"• {s.lstrip('•- ').strip()}", bull_sty))
        else:
            story.append(Paragraph(s, body_sty))

    story.append(Paragraph(f"◆ TailorMyCV Executive Edition — {mode}", brand_sty))
    doc_rl.build(story)
    buf.seek(0)
    return buf.read()


# ════════════════════════════════════════════════════════════
#  RAZORPAY LINK
# ════════════════════════════════════════════════════════════
def get_razorpay_link() -> str:
    try:
        return st.secrets["RAZORPAY_PAYMENT_LINK"]
    except Exception:
        return os.environ.get("RAZORPAY_PAYMENT_LINK", "https://rzp.io/l/tailormycv-demo")


# ════════════════════════════════════════════════════════════
#  UI HELPERS
# ════════════════════════════════════════════════════════════
def render_masthead():
    st.markdown("""
    <div class="masthead">
      <div class="masthead-eyebrow">◆ Executive Edition · V2</div>
      <div class="masthead-logo">Tailor<span class="accent">My</span>CV</div>
      <div class="masthead-sub">AI-Powered Resume Intelligence for Senior Candidates</div>
      <div class="masthead-badges">
        <span class="mbadge">Claude AI</span>
        <span class="mbadge">ATS-Scored</span>
        <span class="mbadge">₹20 / Download</span>
        <span class="mbadge">Instant Results</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

def render_step_bar(current: int):
    steps = ["Upload CV", "Job Description", "Choose Mode", "Result & Pay"]
    pills = ""
    for i, label in enumerate(steps, 1):
        cls  = "done" if i < current else ("active" if i == current else "")
        icon = "✓" if i < current else str(i)
        pills += f'<div class="step-pill {cls}">{icon} {label}</div>'
        if i < len(steps):
            pills += '<div class="step-connector"></div>'
    st.markdown(f'<div class="step-bar">{pills}</div>', unsafe_allow_html=True)

def render_ats_widget(score, color, bar_color, gap):
    st.markdown(f"""
    <div class="ats-score-wrap">
      <div class="ats-score-header">
        <div class="ats-label">ATS Compatibility Score</div>
        <div class="ats-number {color}">{score}%</div>
      </div>
      <div class="ats-bar-track">
        <div class="ats-bar-fill" style="width:{score}%;background:{bar_color};"></div>
      </div>
      <div class="ats-gap"><strong>Gap Analysis:</strong> {gap}</div>
    </div>
    """, unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════
#  STEP PAGES
# ════════════════════════════════════════════════════════════
def page_step1():
    st.markdown('<div class="card"><div class="card-title">📄 Upload or Paste Your CV</div>', unsafe_allow_html=True)
    st.markdown('<div class="info-box">Supports PDF and DOCX uploads, or paste your CV text directly below.</div>', unsafe_allow_html=True)

    uploaded = st.file_uploader("Upload CV (PDF or DOCX)", type=["pdf","docx"], label_visibility="collapsed")
    if uploaded:
        raw = uploaded.read()
        extracted = extract_pdf(raw) if uploaded.name.lower().endswith(".pdf") else extract_docx(raw)
        if extracted.strip():
            st.session_state.cv_text = extracted
            st.success(f"✅ Extracted **{len(extracted.split())}** words from `{uploaded.name}`")
        else:
            st.warning("Auto-extraction failed — please paste your CV below.")

    cv_input = st.text_area(
        "Or paste your full CV text here:",
        value=st.session_state.cv_text,
        height=280,
        placeholder="Paste your complete CV here — Name, contact, work experience, education, skills…",
    )
    st.session_state.cv_text = cv_input
    st.markdown('</div>', unsafe_allow_html=True)

    if st.button("Continue → Paste Job Description", use_container_width=True):
        if len(st.session_state.cv_text.strip()) < 50:
            st.error("Please provide your CV content (minimum 50 characters).")
        else:
            st.session_state.step = 2
            st.rerun()


def page_step2():
    st.markdown('<div class="card"><div class="card-title">🎯 Paste the Job Description</div>', unsafe_allow_html=True)
    st.markdown('<div class="info-box">Copy the full JD from LinkedIn, Naukri, or Foundit. More detail = better tailoring.</div>', unsafe_allow_html=True)

    jd_input = st.text_area(
        "Job Description:",
        value=st.session_state.jd_text,
        height=300,
        placeholder="Paste the complete job description here — responsibilities, required skills, qualifications…",
    )
    st.session_state.jd_text = jd_input
    st.markdown('</div>', unsafe_allow_html=True)

    c1, c2 = st.columns([1, 2])
    with c1:
        if st.button("← Back", use_container_width=True):
            st.session_state.step = 1; st.rerun()
    with c2:
        if st.button("Continue → Select Mode", use_container_width=True):
            if len(st.session_state.jd_text.strip()) < 50:
                st.error("Please paste the job description (minimum 50 characters).")
            else:
                st.session_state.step = 3; st.rerun()


def page_step3():
    st.markdown('<div class="card"><div class="card-title">⚙️ Select Your Tailoring Strategy</div>', unsafe_allow_html=True)

    cols = st.columns(3)
    for idx, (key, meta) in enumerate(MODE_META.items()):
        with cols[idx]:
            selected = st.session_state.mode == key
            st.markdown(f"""
            <div class="mode-card {'selected' if selected else ''}">
              <div class="mode-icon">{meta['icon']}</div>
              <div class="mode-label">{key}</div>
              <div class="mode-desc">{meta['desc']}</div>
            </div>
            """, unsafe_allow_html=True)
            label = "✓ Selected" if selected else "Select"
            if st.button(label, key=f"mode_{key}", use_container_width=True):
                st.session_state.mode = key; st.rerun()

    st.markdown(f"""
    <div class="gold-box" style="margin-top:1rem;">
      ◆ &nbsp;<strong>{st.session_state.mode}</strong> activated — {MODE_META[st.session_state.mode]['desc']}
    </div>
    """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    c1, c2 = st.columns([1, 2])
    with c1:
        if st.button("← Back", use_container_width=True):
            st.session_state.step = 2; st.rerun()
    with c2:
        if st.button("◆ Generate My Executive CV", use_container_width=True):
            with st.spinner("🤖 Claude is engineering your tailored CV…"):
                try:
                    result = call_claude(st.session_state.cv_text, st.session_state.jd_text, st.session_state.mode)
                    st.session_state.tailored_cv   = result["cv_text"]
                    st.session_state.ats_score     = result["score"]
                    st.session_state.ats_color     = result["color"]
                    st.session_state.ats_bar_color = result["bar_color"]
                    st.session_state.gap_analysis  = result["gap"]
                    st.session_state.step          = 4
                    st.rerun()
                except anthropic.AuthenticationError:
                    st.error("❌ Invalid API key. Check Streamlit Secrets → ANTHROPIC_API_KEY.")
                except anthropic.NotFoundError:
                    st.error(f"❌ Model '{CLAUDE_MODEL}' not available on your Anthropic account tier.")
                except Exception as e:
                    st.error(f"❌ Generation failed: {e}")


def page_step4():
    # ATS Score — always visible
    if st.session_state.ats_score is not None:
        render_ats_widget(
            st.session_state.ats_score,
            st.session_state.ats_color,
            st.session_state.ats_bar_color,
            st.session_state.gap_analysis,
        )

    st.markdown(f"""
    <div class="card" style="margin-bottom:0.5rem;">
      <div class="card-title">📋 Your Tailored CV <span class="badge">{st.session_state.mode}</span></div>
    </div>
    """, unsafe_allow_html=True)

    # ── PAYMENT GATE ────────────────────────────────────────
    if not st.session_state.payment_verified:
        preview = st.session_state.tailored_cv[:320] + "\n\n… [Full content locked — unlock below]"
        st.text_area("Preview (locked):", value=preview, height=150, disabled=True, label_visibility="collapsed")

        st.markdown("""
        <div class="warn-box">
          🔒 Your CV is crafted & ATS-scored. Pay <strong>₹20</strong> to unlock the full
          editable text + Word &amp; PDF downloads.
        </div>
        """, unsafe_allow_html=True)

        st.markdown(f"""
        <div class="pay-btn-wrap">
          <a href="{get_razorpay_link()}" target="_blank" class="pay-btn">
            💳 &nbsp; Pay ₹20 — Unlock Full CV + Downloads
          </a>
        </div>
        <p style="text-align:center;font-size:0.76rem;color:#566D84;margin-top:0.6rem;">
          Secure payment via Razorpay · UPI · Cards · Net Banking · Wallets
        </p>
        """, unsafe_allow_html=True)

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown("**Already paid?** Tap below to confirm and unlock your files.")
        _, mid, _ = st.columns([1, 2, 1])
        with mid:
            if st.button("✅ I've Paid — Unlock Now", use_container_width=True):
                st.session_state.payment_verified = True; st.rerun()

    else:
        # ── UNLOCKED ──────────────────────────────────────
        st.markdown('<div class="success-box">✅ Payment confirmed — full CV unlocked. Edit below, then download.</div>', unsafe_allow_html=True)

        edited = st.text_area("Edit your CV (optional fine-tuning):", value=st.session_state.tailored_cv, height=440)
        st.session_state.tailored_cv = edited

        st.markdown(
            '<span class="badge">✓ ATS-Optimised</span>'
            '<span class="badge">✓ Action Verbs</span>'
            '<span class="badge">✓ Quantified Impact</span>'
            '<span class="badge">✓ Executive Format</span>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

        # PDF warning
        st.markdown("""
        <div class="warn-box">
          ⚠️ <strong>ATS Upload Tip:</strong> For older ATS portals (Taleo, legacy Workday),
          use the <strong>Word (.docx)</strong> file for 100% parse accuracy.
          Use PDF for modern portals or direct email applications.
        </div>
        """, unsafe_allow_html=True)

        # Download buttons
        c_docx, c_pdf = st.columns(2)
        with c_docx:
            docx_bytes = generate_docx(st.session_state.tailored_cv, st.session_state.mode)
            if st.download_button(
                "⬇️  Download Word (.docx)", data=docx_bytes,
                file_name=f"TailorMyCV_{st.session_state.mode.replace(' ','-')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            ):
                st.session_state.downloaded = True

        with c_pdf:
            if REPORTLAB_OK:
                pdf_bytes = generate_pdf(st.session_state.tailored_cv, st.session_state.mode)
                if st.download_button(
                    "⬇️  Download PDF", data=pdf_bytes,
                    file_name=f"TailorMyCV_{st.session_state.mode.replace(' ','-')}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                ):
                    st.session_state.downloaded = True
            else:
                st.markdown('<div class="info-box">Add <code>reportlab</code> to requirements.txt to enable PDF export.</div>', unsafe_allow_html=True)

        # Morale booster after download
        if st.session_state.downloaded:
            if st.session_state.morale_msg is None:
                st.session_state.morale_msg = random.choice(MORALE_MESSAGES)
            icon, msg = st.session_state.morale_msg
            st.markdown(f"""
            <div class="morale-card">
              <div class="morale-icon">{icon}</div>
              <div class="morale-text">"{msg}"</div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown("""
            <div class="referral-box">
              Know someone stuck in the ATS grind? 🤝<br>
              <strong>Share TailorMyCV</strong> with friends, batchmates, or your LinkedIn network —
              help them beat the bots for just ₹20.
            </div>
            """, unsafe_allow_html=True)

    # Navigation
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    c1, c2 = st.columns([1, 2])
    with c1:
        if st.button("← Regenerate", use_container_width=True):
            st.session_state.step             = 3
            st.session_state.payment_verified = False
            st.session_state.tailored_cv      = ""
            st.session_state.downloaded       = False
            st.session_state.morale_msg       = None
            st.rerun()
    with c2:
        if st.button("◆ Start Fresh (New CV)", use_container_width=True):
            for k, v in DEFAULTS.items():
                st.session_state[k] = v
            st.rerun()


# ════════════════════════════════════════════════════════════
#  MAIN ROUTER
# ════════════════════════════════════════════════════════════
def main():
    render_masthead()
    render_step_bar(st.session_state.step)
    if   st.session_state.step == 1: page_step1()
    elif st.session_state.step == 2: page_step2()
    elif st.session_state.step == 3: page_step3()
    elif st.session_state.step == 4: page_step4()

    st.markdown("""
    <div class="footer">
      Made with ◆ in India &nbsp;·&nbsp; TailorMyCV Executive Edition © 2025 &nbsp;·&nbsp;
      <a href="mailto:support@tailormycv.in">support@tailormycv.in</a>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
